"""Generate .docx plans and reports — merge reporter.py + fill_weekly_report.py.

Supports two modes:
1. Generate from scratch (plan/report data → new .docx)
2. Fill existing template (owner sends .docx → fill marks column)

CLI usage:
    python -m services.docx_generator plan --data plan.json --output plan.docx
    python -m services.docx_generator report --data report.json --output report.docx
    python -m services.docx_generator fill --template template.docx --data report.json --output filled.docx
    python -m services.docx_generator validate --doc report.docx
"""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from config.settings import load_config, output_error, output_json


# ---------------------------------------------------------------------------
# Cyrillic font helpers (from task-control/reporter.py)
# ---------------------------------------------------------------------------

def _set_cyrillic_fonts(run, font_name: str = "Times New Roman") -> None:
    """Set rFonts for Cyrillic support on run XML."""
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:cs"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)


def _apply_run_format(
    run, font_name: str = "Times New Roman", font_size_pt: int = 12, bold: bool = False
) -> None:
    """Apply font, size, bold, and Cyrillic fonts to a run."""
    run.bold = bold
    run.font.size = Pt(font_size_pt)
    _set_cyrillic_fonts(run, font_name)


# ---------------------------------------------------------------------------
# Table helpers (from task-control/reporter.py)
# ---------------------------------------------------------------------------

def _set_cell_text(
    cell, text: str, font_name: str = "Times New Roman", font_size_pt: int = 12,
    bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Set cell text with formatting, clearing junk indents."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    para = cell.paragraphs[0]
    para.alignment = alignment
    pf = para.paragraph_format
    pf.first_line_indent = None
    pf.left_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for r in para.runs:
        r._element.getparent().remove(r._element)

    # Handle multi-line text
    lines = text.split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i > 0:
            run = para.add_run("\n")
            _apply_run_format(run, font_name, font_size_pt, bold)
        run = para.add_run(line.lstrip("\t"))
        _apply_run_format(run, font_name, font_size_pt, bold)


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0) -> None:
    """Set cell margins in dxa (1/20 of a point)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")
        old = tcMar.find(qn(f"w:{side}"))
        if old is not None:
            tcMar.remove(old)
        tcMar.append(elem)


def _set_table_grid(table, col_widths_twips: list[int]) -> None:
    """Set precise column widths via tblGrid."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    total = sum(col_widths_twips)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is not None:
        tblPr.remove(layout)

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)

    grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    tblPr_idx = list(tbl).index(tblPr)
    tbl.insert(tblPr_idx + 1, grid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_twips):
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(col_widths_twips[i]))
                tcW.set(qn("w:type"), "dxa")


def _page_width_twips(config: dict) -> int:
    """Calculate usable page width in twips."""
    page_cfg = config.get("report", {}).get("page", {})
    total_cm = page_cfg.get("width_cm", 29.7)
    left_cm = page_cfg.get("margin_left_cm", 3.0)
    right_cm = page_cfg.get("margin_right_cm", 1.0)
    usable_cm = total_cm - left_cm - right_cm
    return int(usable_cm * 567)  # 1 cm ~ 567 twips


def _add_merged_row(
    table, text: str, font_name: str, font_size_pt: int, bold: bool = True,
) -> None:
    """Add a horizontally merged row spanning all 5 columns."""
    row = table.add_row()
    cell = row.cells[0]
    for i in range(1, 5):
        cell.merge(row.cells[i])

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _apply_run_format(run, font_name, font_size_pt, bold)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Page setup (landscape for weekly-ops)
# ---------------------------------------------------------------------------

def _setup_page(doc: Document, config: dict) -> None:
    """Set landscape page dimensions and margins."""
    page_cfg = config.get("report", {}).get("page", {})
    section = doc.sections[0]
    # Landscape: width > height
    section.page_width = Cm(page_cfg.get("width_cm", 29.7))
    section.page_height = Cm(page_cfg.get("height_cm", 21.0))
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Cm(page_cfg.get("margin_left_cm", 3.0))
    section.right_margin = Cm(page_cfg.get("margin_right_cm", 1.0))
    section.top_margin = Cm(page_cfg.get("margin_top_cm", 2.0))
    section.bottom_margin = Cm(page_cfg.get("margin_bottom_cm", 2.0))


# ---------------------------------------------------------------------------
# Mode 1: Generate from scratch
# ---------------------------------------------------------------------------

def generate_plan(data: dict, output_path: str, config: dict | None = None) -> str:
    """Generate weekly/monthly plan .docx from scratch."""
    cfg = config or load_config()
    report_cfg = cfg.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 12)
    title_font_size = report_cfg.get("title_font_size_pt", 14)

    doc = Document()
    _setup_page(doc, cfg)

    # Approval header
    _create_approval_header(doc, data, cfg)

    # Title
    title = data.get("title", "ПЛАН мероприятий")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    _apply_run_format(run, font_name, title_font_size, bold=True)
    doc.add_paragraph()

    # Table
    items = data.get("items", [])
    _create_data_table(doc, items, cfg, font_name, font_size)

    # Signature
    _create_signature_block(doc, data, cfg)

    doc.save(output_path)
    return output_path


def generate_report(data: dict, output_path: str, config: dict | None = None) -> str:
    """Generate weekly/monthly report .docx from scratch.

    data = {
        planned: [{item_number, description, deadline, responsible, completion_note}],
        unplanned: [...],
        approver_name, approver_position, signer_name, signer_position,
        title, period_start, period_end
    }
    """
    cfg = config or load_config()
    report_cfg = cfg.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 12)
    title_font_size = report_cfg.get("title_font_size_pt", 14)

    doc = Document()
    _setup_page(doc, cfg)

    # Approval header
    _create_approval_header(doc, data, cfg)

    # Title
    title = data.get("title", "ОТЧЁТ")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    _apply_run_format(run, font_name, title_font_size, bold=True)
    doc.add_paragraph()

    # Planned section
    planned = data.get("planned", [])
    unplanned = data.get("unplanned", [])

    _create_report_table(doc, planned, unplanned, cfg, font_name, font_size)

    # Signature
    _create_signature_block(doc, data, cfg)

    doc.save(output_path)
    return output_path


def _create_data_table(
    doc: Document, items: list[dict], config: dict,
    font_name: str, font_size: int,
) -> None:
    """Create 5-column data table for plan."""
    columns = config.get("report", {}).get("columns", [
        "№ п/п", "Мероприятия", "Сроки проведения", "Ответственный", "Отметка о выполнении"
    ])

    table = doc.add_table(rows=1 + len(items), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    total_w = _page_width_twips(config)
    pcts = config.get("report", {}).get("col_widths_pct", [6, 38, 16, 20, 20])
    col_widths = [int(total_w * p / 100) for p in pcts]
    _set_table_grid(table, col_widths)

    # Header row
    for i, col_name in enumerate(columns):
        _set_cell_text(
            table.rows[0].cells[i], col_name, font_name, font_size,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_margins(table.rows[0].cells[i], 28, 28, 57, 57)

    # Data rows
    for row_idx, item in enumerate(items, 1):
        num = str(item.get("item_number", row_idx))
        desc = item.get("description", "")
        deadline = item.get("deadline", "")
        responsible = item.get("responsible", "")
        if isinstance(responsible, list):
            responsible = ", ".join(
                r.get("value", str(r)) if isinstance(r, dict) else str(r)
                for r in responsible
            )
        completion = item.get("completion_note", "")

        row_data = [num, desc, deadline, str(responsible), completion]
        for i, text in enumerate(row_data):
            _set_cell_text(table.rows[row_idx].cells[i], text, font_name, font_size)
            _set_cell_margins(table.rows[row_idx].cells[i], 28, 28, 57, 57)


def _create_report_table(
    doc: Document, planned: list[dict], unplanned: list[dict], config: dict,
    font_name: str, font_size: int,
) -> None:
    """Create report table with planned + unplanned sections."""
    columns = config.get("report", {}).get("columns", [
        "№ п/п", "Мероприятия", "Сроки проведения", "Ответственный", "Отметка о выполнении"
    ])

    # Count rows: header + "Планируемые" merged + planned + "Доп" merged + unplanned
    n_planned = len(planned)
    n_unplanned = len(unplanned)
    total_rows = 1  # header
    total_rows += 1  # "Планируемые мероприятия" merged row
    total_rows += n_planned
    if n_unplanned:
        total_rows += 1  # "Дополнительные мероприятия" merged row
        total_rows += n_unplanned

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    total_w = _page_width_twips(config)
    pcts = config.get("report", {}).get("col_widths_pct", [6, 38, 16, 20, 20])
    col_widths = [int(total_w * p / 100) for p in pcts]
    _set_table_grid(table, col_widths)

    # Header row
    for i, col_name in enumerate(columns):
        _set_cell_text(
            table.rows[0].cells[i], col_name, font_name, font_size,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_margins(table.rows[0].cells[i], 28, 28, 57, 57)

    # "Планируемые мероприятия" section header
    _add_merged_row(table, "Планируемые мероприятия", font_name, font_size, bold=True)

    # Planned rows
    for item in planned:
        _add_item_row(table, item, font_name, font_size)

    # Unplanned section
    if n_unplanned:
        _add_merged_row(table, "Дополнительные мероприятия", font_name, font_size, bold=True)
        for item in unplanned:
            _add_item_row(table, item, font_name, font_size)


def _add_item_row(table, item: dict, font_name: str, font_size: int) -> None:
    """Add a single data row to the table."""
    row = table.add_row()
    num = str(item.get("item_number", ""))
    desc = item.get("description", "")
    deadline = item.get("deadline", "")
    responsible = item.get("responsible", "")
    if isinstance(responsible, list):
        responsible = ", ".join(
            r.get("value", str(r)) if isinstance(r, dict) else str(r)
            for r in responsible
        )
    completion = item.get("completion_note", "")

    values = [num, desc, deadline, str(responsible), completion]
    for i, text in enumerate(values):
        _set_cell_text(row.cells[i], text, font_name, font_size)
        _set_cell_margins(row.cells[i], 28, 28, 57, 57)


# ---------------------------------------------------------------------------
# Mode 2: Fill existing template (from fill_weekly_report.py)
# ---------------------------------------------------------------------------

def load_template(path: str) -> Document:
    """Load .docx template and verify structure."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    doc = Document(path)

    section = doc.sections[0]
    if section.orientation != WD_ORIENT.LANDSCAPE:
        raise ValueError("Expected landscape orientation")

    if len(doc.tables) < 2:
        raise ValueError(f"Expected at least 2 tables, found {len(doc.tables)}")

    return doc


def detect_font_settings(doc: Document) -> dict:
    """Detect font settings from existing data cells."""
    data_table = doc.tables[1]  # DATA_TABLE_IDX = 1
    font_names: list[str] = []
    font_sizes: list[int] = []

    for row in data_table.rows[2:]:  # DATA_ROW_START = 2
        for col_idx in (1, 2, 3):  # COL_ACTIVITY, COL_DATES, COL_RESPONSIBLE
            cell = row.cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.name:
                        font_names.append(run.font.name)
                    if run.font.size is not None:
                        font_sizes.append(run.font.size)

    detected_name = "Times New Roman"
    detected_size = Pt(12)

    if font_names:
        detected_name = max(set(font_names), key=font_names.count)
    if font_sizes:
        detected_size = max(set(font_sizes), key=font_sizes.count)

    return {"font_name": detected_name, "font_size": detected_size}


def fill_template(
    doc: Document,
    report_data: dict,
    font_settings: dict,
) -> Document:
    """Fill marks column in existing template."""
    data_table = doc.tables[1]
    font_name = font_settings["font_name"]
    font_size = font_settings["font_size"]

    planned = report_data.get("planned", [])
    unplanned = report_data.get("unplanned", [])

    # Fill marks for planned items (match by activity text)
    for row in data_table.rows[2:]:
        activity_text = row.cells[1].text.strip()
        if not activity_text:
            continue

        # Find matching report item
        mark_text = _find_mark_for_activity(activity_text, planned)
        if mark_text:
            _set_template_cell(row.cells[4], mark_text, font_name, font_size)

    # Add unplanned section if any
    if unplanned:
        done_unplanned = [u for u in unplanned if u.get("status", "done") == "done"
                          or u.get("completion_note")]
        if done_unplanned:
            _add_merged_row(data_table, "Дополнительные мероприятия",
                            font_name, font_size // 12700 if font_size > 100 else 12, bold=True)
            for item in done_unplanned:
                row = data_table.add_row()
                num = str(item.get("item_number", ""))
                desc = item.get("description", "")
                deadline = item.get("deadline", "")
                responsible = item.get("responsible", "")
                mark = item.get("completion_note", "")

                values = [num, desc, deadline, str(responsible), mark]
                for i, text in enumerate(values):
                    _set_template_cell(row.cells[i], text, font_name, font_size)

    return doc


def _find_mark_for_activity(activity: str, items: list[dict]) -> str | None:
    """Find completion note for an activity by keyword overlap."""
    activity_lower = activity.lower()
    activity_words = {w for w in activity_lower.split() if len(w) > 3}

    for item in items:
        desc = (item.get("description") or "").lower()

        # Exact substring match
        if desc and desc in activity_lower:
            return item.get("completion_note", "")

        # Keyword overlap
        desc_words = {w for w in desc.split() if len(w) > 3}
        if activity_words and desc_words:
            overlap = activity_words & desc_words
            if len(overlap) >= 2:
                return item.get("completion_note", "")

    return None


def _set_template_cell(cell, text: str, font_name: str, font_size) -> None:
    """Set cell text in template, preserving style compatibility."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    p = cell.paragraphs[0]
    lines = text.split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = font_size


# ---------------------------------------------------------------------------
# Validation (from fill_weekly_report.py)
# ---------------------------------------------------------------------------

def validate_formatting(doc: Document) -> list[str]:
    """Validate column 4 formatting and completeness."""
    issues: list[str] = []
    data_table = doc.tables[1]
    font_settings = detect_font_settings(doc)
    expected_name = font_settings["font_name"]
    expected_size = font_settings["font_size"]

    for ri, row in enumerate(data_table.rows[2:], start=2):
        cell = row.cells[4]
        activity_cell = row.cells[1]
        activity_text = activity_cell.text.strip()

        # Skip merged/section rows
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None and int(gs.get(qn("w:val"), "1")) > 1:
                continue

        mark_text = cell.text.strip()

        if activity_text and not mark_text:
            issues.append(f"R{ri}: пустая отметка для «{activity_text[:50]}»")
            continue

        if not mark_text:
            continue

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                actual_name = run.font.name
                if actual_name and actual_name != expected_name:
                    issues.append(f"R{ri}/C4: шрифт «{actual_name}» вместо «{expected_name}»")
                actual_size = run.font.size
                if actual_size is not None and actual_size != expected_size:
                    expected_pt = expected_size // 12700
                    actual_pt = actual_size // 12700
                    issues.append(f"R{ri}/C4: размер {actual_pt}pt вместо {expected_pt}pt")

    return issues


# ---------------------------------------------------------------------------
# Common blocks
# ---------------------------------------------------------------------------

def _create_approval_header(doc: Document, data: dict, config: dict) -> None:
    """Create УТВЕРЖДАЮ approval header aligned to the right."""
    report_cfg = config.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    title_font_size = report_cfg.get("title_font_size_pt", 14)

    approver_name = data.get("approver_name", "")
    approver_position = data.get("approver_position", "")

    if not approver_name:
        return

    header_text = report_cfg.get("approval_header", "УТВЕРЖДАЮ")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(header_text)
    _apply_run_format(run, font_name, title_font_size, bold=True)

    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = para2.add_run(approver_position)
    _apply_run_format(run2, font_name, title_font_size)

    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = para3.add_run(f"__________ {approver_name}")
    _apply_run_format(run3, font_name, title_font_size)

    para4 = doc.add_paragraph()
    para4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run4 = para4.add_run(f'"____" ____________ {date.today().year} г.')
    _apply_run_format(run4, font_name, title_font_size)

    doc.add_paragraph()


def _create_signature_block(doc: Document, data: dict, config: dict) -> None:
    """Create signature block after the table."""
    report_cfg = config.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    title_font_size = report_cfg.get("title_font_size_pt", 14)

    position = data.get("signer_position", "")
    name = data.get("signer_name", "")

    if not position and not name:
        return

    doc.add_paragraph()
    doc.add_paragraph()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(f"{position}{'  ' * 10}__________ {name}")
    _apply_run_format(run, font_name, title_font_size)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="DOCX generator for weekly-ops")
    sub = parser.add_subparsers(dest="command", required=True)

    p_plan = sub.add_parser("plan", help="Generate plan .docx")
    p_plan.add_argument("--data", required=True, help="Plan data JSON")
    p_plan.add_argument("--output", required=True, help="Output .docx path")

    p_report = sub.add_parser("report", help="Generate report .docx")
    p_report.add_argument("--data", required=True, help="Report data JSON")
    p_report.add_argument("--output", required=True, help="Output .docx path")

    p_fill = sub.add_parser("fill", help="Fill existing template")
    p_fill.add_argument("--template", required=True, help="Template .docx path")
    p_fill.add_argument("--data", required=True, help="Report data JSON")
    p_fill.add_argument("--output", required=True, help="Output .docx path")

    p_val = sub.add_parser("validate", help="Validate document formatting")
    p_val.add_argument("--doc", required=True, help="Document .docx path")

    args = parser.parse_args()
    config = load_config()

    try:
        if args.command == "plan":
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
            result = generate_plan(data, args.output, config)
            output_json({"path": result, "success": True})

        elif args.command == "report":
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
            result = generate_report(data, args.output, config)
            output_json({"path": result, "success": True})

        elif args.command == "fill":
            doc = load_template(args.template)
            with open(args.data, "r", encoding="utf-8") as f:
                report_data = json.load(f)
            font_settings = detect_font_settings(doc)
            fill_template(doc, report_data, font_settings)
            issues = validate_formatting(doc)
            doc.save(args.output)
            output_json({
                "path": args.output,
                "success": True,
                "issues": issues,
            })

        elif args.command == "validate":
            doc = load_template(args.doc)
            issues = validate_formatting(doc)
            output_json({"issues": issues, "count": len(issues)})

    except (RuntimeError, ValueError, FileNotFoundError) as e:
        output_error(str(e))


if __name__ == "__main__":
    main()
