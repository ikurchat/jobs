"""Tests for .docx report generation."""

import json
import os
from pathlib import Path

import pytest
from docx import Document

os.environ.setdefault("BASEROW_URL", "https://baserow.example.com")
os.environ.setdefault("BASEROW_TOKEN", "test-token-123")

from services.reporter import (
    create_discipline_report,
    create_monthly_plan,
    create_weekly_plan,
    create_weekly_report,
    create_weekly_report_with_unplanned,
)


@pytest.fixture
def plan_data():
    return {
        "period_start": "03.02.2025",
        "period_end": "09.02.2025",
        "approver_name": "Иванов И.И.",
        "approver_position": "Директор по безопасности",
        "signer_name": "Смирнов И.Ю.",
        "signer_position": "Начальник отдела ИБ",
        "title": "ПЛАН мероприятий ЦОК на 03.02–09.02.2025",
        "items": [
            {
                "item_number": 1,
                "description": "Мониторинг инцидентов ИБ",
                "deadline": "03.02–09.02",
                "responsible": "Все",
            },
            {
                "item_number": 2,
                "description": "Ретроспективный анализ",
                "deadline": "03.02–09.02",
                "responsible": "Иванов И.И., Петров П.П.",
            },
            {
                "item_number": 3,
                "description": "Совещание с Интегратор",
                "deadline": "10.02",
                "responsible": "Смирнов И.Ю.",
            },
        ],
    }


@pytest.fixture
def report_data(plan_data):
    data = dict(plan_data)
    data["title"] = "ОТЧЁТ ЦОК за 03.02–09.02.2025"
    for item in data["items"]:
        item["completion_note"] = "Выполнено."
    return data


@pytest.fixture
def unplanned_data():
    return {
        "title": "Дополнительные мероприятия (внеплановые)",
        "items": [
            {
                "item_number": 1,
                "description": "Справка по Васильеву",
                "deadline": "09.02–13.02",
                "responsible": "Смирнов И.Ю.",
                "completion_note": "Выполнено.",
            },
            {
                "item_number": 2,
                "description": "Изучение договора с Вендором-SOAR",
                "deadline": "09.02–13.02",
                "responsible": "Смирнов И.Ю.",
                "completion_note": "В работе.",
            },
        ],
    }


class TestWeeklyPlan:
    def test_creates_valid_docx(self, plan_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "plan.docx")
        result = create_weekly_plan(plan_data, output, config)

        assert Path(result).exists()
        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_has_5_columns(self, plan_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "plan.docx")
        create_weekly_plan(plan_data, output, config)

        doc = Document(output)
        tables = doc.tables
        assert len(tables) >= 1
        # First table should have 5 columns
        assert len(tables[0].columns) == 5

    def test_all_items_present(self, plan_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "plan.docx")
        create_weekly_plan(plan_data, output, config)

        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        assert "Мониторинг инцидентов ИБ" in table_text
        assert "Ретроспективный анализ" in table_text
        assert "Совещание с Интегратор" in table_text

    def test_has_approval_header(self, plan_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "plan.docx")
        create_weekly_plan(plan_data, output, config)

        doc = Document(output)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "УТВЕРЖДАЮ" in all_text


class TestWeeklyReport:
    def test_creates_valid_docx(self, report_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "report.docx")
        result = create_weekly_report(report_data, output, config)
        assert Path(result).exists()

    def test_completion_notes(self, report_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "report.docx")
        create_weekly_report(report_data, output, config)

        doc = Document(output)
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        assert "Выполнено" in table_text


class TestReportWithUnplanned:
    def test_two_tables(self, plan_data, unplanned_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "report_full.docx")
        create_weekly_report_with_unplanned(plan_data, unplanned_data, output, config)

        doc = Document(output)
        tables = doc.tables
        assert len(tables) >= 2

    def test_unplanned_items_present(self, plan_data, unplanned_data, tmp_work_dir, config):
        output = str(tmp_work_dir / "report_full.docx")
        create_weekly_report_with_unplanned(plan_data, unplanned_data, output, config)

        doc = Document(output)
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "

        assert "Справка по Васильеву" in all_text
        assert "Вендором-SOAR" in all_text


class TestEmptyReport:
    def test_empty_items(self, tmp_work_dir, config):
        data = {
            "approver_name": "",
            "approver_position": "",
            "signer_name": "",
            "signer_position": "",
            "title": "Empty Report",
            "items": [],
        }
        output = str(tmp_work_dir / "empty.docx")
        result = create_weekly_report(data, output, config)
        assert Path(result).exists()

        doc = Document(result)
        # Should still have a table with header row
        assert len(doc.tables) >= 1


class TestMonthlyPlan:
    def test_creates_valid_docx(self, tmp_work_dir, config):
        data = {
            "month": "Февраль 2025",
            "approver_name": "Директор",
            "approver_position": "Директор",
            "signer_name": "Начальник",
            "signer_position": "Начальник отдела",
            "title": "ПЛАН мероприятий на февраль 2025",
            "items": [
                {"item_number": 1, "description": "Мониторинг", "deadline": "Февраль", "responsible": "Все"},
            ],
        }
        output = str(tmp_work_dir / "monthly.docx")
        result = create_monthly_plan(data, output, config)
        assert Path(result).exists()


class TestDisciplineReport:
    def test_creates_valid_docx(self, tmp_work_dir, config):
        data = {
            "title": "Отчёт по дисциплине",
            "period_start": "03.02.2025",
            "period_end": "09.02.2025",
            "employees": [
                {"fio": "Иванов И.И.", "tasks_total": 8, "tasks_completed": 6, "tasks_overdue": 2, "on_time_rate": 75},
                {"fio": "Петров П.П.", "tasks_total": 5, "tasks_completed": 5, "tasks_overdue": 0, "on_time_rate": 100},
            ],
        }
        output = str(tmp_work_dir / "discipline.docx")
        result = create_discipline_report(data, output, config)
        assert Path(result).exists()

        doc = Document(result)
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        assert "Иванов" in all_text
        assert "Петров" in all_text
