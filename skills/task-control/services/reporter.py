"""Report generator — .docx plans and reports (weekly, monthly, discipline).

Following doc-review patterns: python-docx, cyrillic fonts, table formatting.

CLI usage:
    python -m services.reporter weekly_plan --input data.json --output plan.docx
    python -m services.reporter weekly_report --input data.json --output report.docx
    python -m services.reporter weekly_report_with_unplanned --plan plan.json --unplanned unplanned.json --output report.docx
    python -m services.reporter monthly_plan --input data.json --output plan.docx
    python -m services.reporter discipline_report --input data.json --output discipline.docx
"""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

from config.settings import load_config, output_error, output_json


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cyrillic_fonts(run, font_name: str = "Times New Roman") -> None:
    """Set rFonts for Cyrillic support on run XML."""
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:cs"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)


def _apply_run_format(
    run, font_name: str = "Times New Roman", font_size_pt: int = 14, bold: bool = False
) -> None:
    """Apply font, size, bold, and Cyrillic fonts to a run."""
    run.bold = bold
    run.font.size = Pt(font_size_pt)
    _set_cyrillic_fonts(run, font_name)


def _setup_page(doc: Document, config: dict) -> None:
    """Set page dimensions and margins."""
    page_cfg = config.get("report", {}).get("page", {})
    section = doc.sections[0]
    section.page_width = Cm(page_cfg.get("width_cm", 21.0))
    section.page_height = Cm(page_cfg.get("height_cm", 29.7))
    section.left_margin = Cm(page_cfg.get("margin_left_cm", 3.0))
    section.right_margin = Cm(page_cfg.get("margin_right_cm", 1.0))
    section.top_margin = Cm(page_cfg.get("margin_top_cm", 2.0))
    section.bottom_margin = Cm(page_cfg.get("margin_bottom_cm", 2.0))


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0) -> None:
    """Set cell margins in dxa (1/20 of a point)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")
        old = tcMar.find(qn(f"w:{side}"))
        if old is not None:
            tcMar.remove(old)
        tcMar.append(elem)


def _set_cell_text(
    cell, text: str, font_name: str = "Times New Roman", font_size_pt: int = 14,
    bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT
) -> None:
    """Set cell text with formatting, clearing junk indents."""
    # Clear extra paragraphs
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    para = cell.paragraphs[0]
    para.alignment = alignment
    # Clear junk indents from default styles
    pf = para.paragraph_format
    pf.first_line_indent = None
    pf.left_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Clear existing runs
    for r in para.runs:
        r._element.getparent().remove(r._element)

    run = para.add_run(text.lstrip("\t"))
    _apply_run_format(run, font_name, font_size_pt, bold)


def _page_width_twips(config: dict) -> int:
    """Calculate usable page width in twips."""
    page_cfg = config.get("report", {}).get("page", {})
    total_cm = page_cfg.get("width_cm", 21.0)
    left_cm = page_cfg.get("margin_left_cm", 3.0)
    right_cm = page_cfg.get("margin_right_cm", 1.0)
    usable_cm = total_cm - left_cm - right_cm
    return int(usable_cm * 567)  # 1 cm ≈ 567 twips


def _set_table_grid(table, col_widths_twips: list[int]) -> None:
    """Set precise column widths via tblGrid."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    total = sum(col_widths_twips)

    # Set total width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")

    # Remove fixed layout
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is not None:
        tblPr.remove(layout)

    # Create tblGrid
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)

    grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    # Insert grid right after tblPr
    tblPr_idx = list(tbl).index(tblPr)
    tbl.insert(tblPr_idx + 1, grid)

    # Set cell widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_twips):
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(col_widths_twips[i]))
                tcW.set(qn("w:type"), "dxa")


def _create_approval_header(
    doc: Document, approver_name: str, approver_position: str, config: dict
) -> None:
    """Create УТВЕРЖДАЮ approval header aligned to the right."""
    report_cfg = config.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 14)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(report_cfg.get("approval_header", "УТВЕРЖДАЮ"))
    _apply_run_format(run, font_name, font_size, bold=True)

    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = para2.add_run(approver_position)
    _apply_run_format(run2, font_name, font_size)

    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = para3.add_run(f"__________ {approver_name}")
    _apply_run_format(run3, font_name, font_size)

    para4 = doc.add_paragraph()
    para4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run4 = para4.add_run(f'"____" ____________ {date.today().year} г.')
    _apply_run_format(run4, font_name, font_size)

    # Empty line
    doc.add_paragraph()


def _create_5col_table(
    doc: Document,
    items: list[dict],
    config: dict,
    title: str = "",
) -> None:
    """Create 5-column plan/report table."""
    report_cfg = config.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 14)
    columns = report_cfg.get("columns", [
        "№ п/п", "Мероприятия", "Сроки проведения", "Ответственный", "Отметка о выполнении"
    ])

    if title:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        _apply_run_format(run, font_name, font_size, bold=True)
        doc.add_paragraph()

    # Create table
    table = doc.add_table(rows=1 + len(items), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Column widths: №(5%), Мероприятия(40%), Сроки(15%), Ответственный(20%), Отметка(20%)
    total_w = _page_width_twips(config)
    col_widths = [
        int(total_w * 0.06),
        int(total_w * 0.38),
        int(total_w * 0.16),
        int(total_w * 0.20),
        int(total_w * 0.20),
    ]
    _set_table_grid(table, col_widths)

    # Header row
    for i, col_name in enumerate(columns):
        _set_cell_text(
            table.rows[0].cells[i], col_name, font_name, font_size,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_margins(table.rows[0].cells[i], 28, 28, 57, 57)

    # Data rows
    for row_idx, item in enumerate(items, 1):
        num = str(item.get("item_number", row_idx))
        desc = item.get("description", "")
        deadline = item.get("deadline", "")
        responsible = item.get("responsible", "")
        if isinstance(responsible, list):
            responsible = ", ".join(
                r.get("value", str(r)) if isinstance(r, dict) else str(r)
                for r in responsible
            )
        completion = item.get("completion_note", "")

        row_data = [num, desc, deadline, str(responsible), completion]
        for i, text in enumerate(row_data):
            _set_cell_text(table.rows[row_idx].cells[i], text, font_name, font_size - 2)
            _set_cell_margins(table.rows[row_idx].cells[i], 28, 28, 57, 57)


def _create_signature_block(
    doc: Document, position: str, name: str, config: dict
) -> None:
    """Create signature block after the table."""
    report_cfg = config.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(f"{position}{'  ' * 10}__________ {name}")
    _apply_run_format(run, font_name, font_size)


# ---------------------------------------------------------------------------
# Document generation functions
# ---------------------------------------------------------------------------

def create_weekly_plan(data: dict, output_path: str, config: dict | None = None) -> str:
    """Generate weekly plan .docx.

    data = {
        period_start, period_end, approver_name, approver_position,
        title, items: [{item_number, description, deadline, responsible}]
    }
    """
    cfg = config or load_config()
    doc = Document()
    _setup_page(doc, cfg)

    # Approval header
    _create_approval_header(
        doc,
        data.get("approver_name", ""),
        data.get("approver_position", ""),
        cfg,
    )

    # Title
    title = data.get("title", f"ПЛАН мероприятий на {data.get('period_start', '')} — {data.get('period_end', '')}")

    # Table
    _create_5col_table(doc, data.get("items", []), cfg, title=title)

    # Signature
    _create_signature_block(
        doc,
        data.get("signer_position", ""),
        data.get("signer_name", ""),
        cfg,
    )

    doc.save(output_path)
    return output_path


def create_weekly_report(data: dict, output_path: str, config: dict | None = None) -> str:
    """Generate weekly report .docx (same format as plan, filled status column)."""
    cfg = config or load_config()
    doc = Document()
    _setup_page(doc, cfg)

    _create_approval_header(
        doc,
        data.get("approver_name", ""),
        data.get("approver_position", ""),
        cfg,
    )

    title = data.get("title", f"ОТЧЁТ за {data.get('period_start', '')} — {data.get('period_end', '')}")
    _create_5col_table(doc, data.get("items", []), cfg, title=title)

    _create_signature_block(
        doc, data.get("signer_position", ""), data.get("signer_name", ""), cfg,
    )

    doc.save(output_path)
    return output_path


def create_weekly_report_with_unplanned(
    plan_data: dict, unplanned_data: dict, output_path: str, config: dict | None = None
) -> str:
    """Generate weekly report with two sections: planned + unplanned."""
    cfg = config or load_config()
    doc = Document()
    _setup_page(doc, cfg)

    _create_approval_header(
        doc,
        plan_data.get("approver_name", ""),
        plan_data.get("approver_position", ""),
        cfg,
    )

    # Section 1: Planned
    _create_5col_table(
        doc,
        plan_data.get("items", []),
        cfg,
        title=plan_data.get("title", "Плановые мероприятия"),
    )

    # Spacer
    report_cfg = cfg.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 14)

    doc.add_paragraph()

    # Section 2: Unplanned
    unplanned_items = unplanned_data.get("items", [])
    if unplanned_items:
        _create_5col_table(
            doc,
            unplanned_items,
            cfg,
            title=unplanned_data.get("title", "Дополнительные мероприятия (внеплановые)"),
        )

    _create_signature_block(
        doc,
        plan_data.get("signer_position", ""),
        plan_data.get("signer_name", ""),
        cfg,
    )

    doc.save(output_path)
    return output_path


def create_monthly_plan(data: dict, output_path: str, config: dict | None = None) -> str:
    """Generate monthly plan .docx."""
    cfg = config or load_config()
    doc = Document()
    _setup_page(doc, cfg)

    _create_approval_header(
        doc,
        data.get("approver_name", ""),
        data.get("approver_position", ""),
        cfg,
    )

    title = data.get("title", f"ПЛАН мероприятий на {data.get('month', '')}")
    _create_5col_table(doc, data.get("items", []), cfg, title=title)

    _create_signature_block(
        doc, data.get("signer_position", ""), data.get("signer_name", ""), cfg,
    )

    doc.save(output_path)
    return output_path


def create_discipline_report(
    data: dict, output_path: str, config: dict | None = None
) -> str:
    """Generate discipline report with per-employee stats."""
    cfg = config or load_config()
    doc = Document()
    _setup_page(doc, cfg)

    report_cfg = cfg.get("report", {})
    font_name = report_cfg.get("font_name", "Times New Roman")
    font_size = report_cfg.get("font_size_pt", 14)

    # Title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(data.get("title", "Отчёт по исполнительской дисциплине"))
    _apply_run_format(run, font_name, font_size, bold=True)
    doc.add_paragraph()

    # Period
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_text = f"Период: {data.get('period_start', '')} — {data.get('period_end', '')}"
    run2 = para2.add_run(period_text)
    _apply_run_format(run2, font_name, font_size)
    doc.add_paragraph()

    # Table: ФИО, Всего, Выполнено, Просрочено, % в срок
    employees = data.get("employees", [])
    table = doc.add_table(rows=1 + len(employees), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    total_w = _page_width_twips(cfg)
    col_widths = [
        int(total_w * 0.30),
        int(total_w * 0.15),
        int(total_w * 0.18),
        int(total_w * 0.18),
        int(total_w * 0.19),
    ]
    _set_table_grid(table, col_widths)

    headers = ["ФИО", "Всего задач", "Выполнено", "Просрочено", "% в срок"]
    for i, h in enumerate(headers):
        _set_cell_text(
            table.rows[0].cells[i], h, font_name, font_size,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_margins(table.rows[0].cells[i], 28, 28, 57, 57)

    for row_idx, emp in enumerate(employees, 1):
        row_data = [
            emp.get("fio", ""),
            str(emp.get("tasks_total", 0)),
            str(emp.get("tasks_completed", 0)),
            str(emp.get("tasks_overdue", 0)),
            f"{emp.get('on_time_rate', 0):.0f}%",
        ]
        for i, text in enumerate(row_data):
            _set_cell_text(table.rows[row_idx].cells[i], text, font_name, font_size - 2)
            _set_cell_margins(table.rows[row_idx].cells[i], 28, 28, 57, 57)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Report generator (.docx)")
    sub = parser.add_subparsers(dest="command", required=True)

    for cmd in ["weekly_plan", "weekly_report", "monthly_plan", "discipline_report"]:
        p = sub.add_parser(cmd)
        p.add_argument("--input", required=True, help="JSON data file")
        p.add_argument("--output", required=True, help="Output .docx path")

    p_wu = sub.add_parser("weekly_report_with_unplanned")
    p_wu.add_argument("--plan", required=True, help="Planned data JSON")
    p_wu.add_argument("--unplanned", required=True, help="Unplanned data JSON")
    p_wu.add_argument("--output", required=True, help="Output .docx path")

    args = parser.parse_args()

    try:
        config = load_config()

        if args.command == "weekly_plan":
            with open(args.input, "r", encoding="utf-8") as f:
                data = json.load(f)
            result = create_weekly_plan(data, args.output, config)
            output_json({"path": result, "success": True})

        elif args.command == "weekly_report":
            with open(args.input, "r", encoding="utf-8") as f:
                data = json.load(f)
            result = create_weekly_report(data, args.output, config)
            output_json({"path": result, "success": True})

        elif args.command == "weekly_report_with_unplanned":
            with open(args.plan, "r", encoding="utf-8") as f:
                plan_data = json.load(f)
            with open(args.unplanned, "r", encoding="utf-8") as f:
                unplanned_data = json.load(f)
            result = create_weekly_report_with_unplanned(plan_data, unplanned_data, args.output, config)
            output_json({"path": result, "success": True})

        elif args.command == "monthly_plan":
            with open(args.input, "r", encoding="utf-8") as f:
                data = json.load(f)
            result = create_monthly_plan(data, args.output, config)
            output_json({"path": result, "success": True})

        elif args.command == "discipline_report":
            with open(args.input, "r", encoding="utf-8") as f:
                data = json.load(f)
            result = create_discipline_report(data, args.output, config)
            output_json({"path": result, "success": True})

    except (RuntimeError, ValueError, json.JSONDecodeError, FileNotFoundError) as e:
        output_error(str(e))


if __name__ == "__main__":
    main()
