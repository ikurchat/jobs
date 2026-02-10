"""Tests for generate_docx.py — document generation."""

import json
import shutil
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils import load_config


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def config():
    return load_config()


@pytest.fixture
def test_dir():
    d = Path(tempfile.mkdtemp())
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture
def sample_content():
    return {
        "title": "Справка о результатах проверки ИБ",
        "addressee": "Директору по безопасности\nИванову Ивану Ивановичу",
        "resume": "В ходе проверки выявлены нарушения политики информационной безопасности.",
        "details": "10 января 2024 года зафиксированы попытки несанкционированного доступа.\n\nИспользовались учётные данные уволенного сотрудника.",
        "conclusions": "Предлагается провести аудит учётных записей и усилить контроль доступа.\n\nСрок исполнения — до 1 марта 2024 года.",
        "appendices": ["Акт проверки от 10.01.2024", "Журнал событий SIEM"],
        "signer_position": "Начальник отдела ИБ",
        "signer_name": "Петров П.П.",
        "executor_name": "Сидоров С.С.",
        "executor_phone": "1234",
    }


# ---------------------------------------------------------------------------
# Tests: document creation
# ---------------------------------------------------------------------------

class TestCreateDocument:
    def test_creates_file(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        result = create_document(sample_content, config, path)
        assert result.exists()
        assert result.suffix == ".docx"

    def test_page_setup(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        section = doc.sections[0]

        # Check margins (with tolerance of 0.1 cm)
        left_cm = section.left_margin / 360000
        right_cm = section.right_margin / 360000
        top_cm = section.top_margin / 360000
        bottom_cm = section.bottom_margin / 360000

        assert abs(left_cm - 3.0) < 0.1, f"Left margin: {left_cm}"
        assert abs(right_cm - 1.0) < 0.1, f"Right margin: {right_cm}"
        assert abs(top_cm - 2.0) < 0.1, f"Top margin: {top_cm}"
        assert abs(bottom_cm - 2.0) < 0.1, f"Bottom margin: {bottom_cm}"

    def test_has_tables(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        # Should have at least: header table, appendix table, signature table
        assert len(doc.tables) >= 3, f"Expected >=3 tables, got {len(doc.tables)}"

    def test_header_table_content(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        header_table = doc.tables[0]

        # Should be 1x2
        assert len(header_table.rows) == 1
        assert len(header_table.columns) == 2

        # Left cell = title
        left_text = header_table.rows[0].cells[0].text
        assert "Справка" in left_text

        # Right cell = addressee
        right_text = header_table.rows[0].cells[1].text
        assert "Иванов" in right_text

    def test_header_table_no_borders(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        header_table = doc.tables[0]

        # Check that borders are set to "none"
        tblPr = header_table._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
        if borders is not None:
            for child in borders:
                val = child.get(qn("w:val"))
                assert val in ("none", "nil", None), f"Border {child.tag} has val={val}"

    def test_header_table_grid(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        header_table = doc.tables[0]
        tbl = header_table._tbl

        # tblGrid must exist with 2 gridCol elements
        grid = tbl.find(qn("w:tblGrid"))
        assert grid is not None, "tblGrid missing"
        cols = grid.findall(qn("w:gridCol"))
        assert len(cols) == 2, f"Expected 2 gridCol, got {len(cols)}"

        # Total width should match page usable width
        page = config["page"]
        expected_twips = int(
            (page["width_cm"] - page["margin_left_cm"] - page["margin_right_cm"])
            / 2.54 * 1440
        )
        actual_total = sum(int(c.get(qn("w:w"))) for c in cols)
        assert abs(actual_total - expected_twips) < 10, \
            f"Grid total {actual_total} != expected {expected_twips}"

        # Left column should be ~60%
        left_w = int(cols[0].get(qn("w:w")))
        assert abs(left_w / actual_total - 0.6) < 0.05, \
            f"Left column ratio {left_w/actual_total:.2f} != ~0.6"

        # No tblLayout fixed
        tblPr = tbl.tblPr
        layout = tblPr.find(qn("w:tblLayout")) if tblPr is not None else None
        if layout is not None:
            assert layout.get(qn("w:type")) != "fixed", "tblLayout should not be fixed"

        # Right cell should be right-aligned
        right_cell = header_table.rows[0].cells[1]
        for p in right_cell.paragraphs:
            if p.text.strip():
                assert p.alignment == 2, f"Right cell alignment: {p.alignment}"

    def test_body_content(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "нарушения политики" in all_text
        assert "несанкционированного доступа" in all_text
        assert "аудит учётных записей" in all_text

    def test_appendix_table(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        # Find table with "Приложение:" text
        appendix_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Приложение" in cell.text:
                        appendix_found = True
                        break
        assert appendix_found, "Appendix table not found"

    def test_signature_table(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        # Last table should be signature
        last_table = doc.tables[-1]
        all_text = ""
        for row in last_table.rows:
            for cell in row.cells:
                all_text += cell.text + " "
        assert "Петров" in all_text or "Начальник" in all_text

    def test_footer_has_executor(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        # Executor is now in document body (NOT footer)
        body_text = " ".join(p.text for p in doc.paragraphs)
        assert "Сидоров" in body_text, \
            "Executor name should appear in document body"
        assert "1234" in body_text, \
            "Executor phone should appear in document body"

    def test_executor_not_in_footer(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        footer = doc.sections[0].footer
        footer_text = " ".join(p.text for p in footer.paragraphs)
        assert "Сидоров" not in footer_text, \
            "Executor name should NOT appear in page footer"

    def test_footer_font_size(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        footer = doc.sections[0].footer
        for p in footer.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.size == Pt(12), \
                        f"Footer font size is {run.font.size}, expected 12pt"

    def test_title_page_flag_for_page_numbering(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        section = doc.sections[0]
        sectPr = section._sectPr
        title_pg = sectPr.find(qn("w:titlePg"))
        assert title_pg is not None, "titlePg flag should be set for page numbering from page 2"

    def test_font_in_body(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        # Executor lines are 12pt, everything else is 14pt
        executor_texts = {sample_content["executor_name"], sample_content["executor_phone"]}
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.name == "Times New Roman", \
                        f"Font is {run.font.name} for text: {run.text[:50]}"
                    if run.text.strip() in executor_texts:
                        assert run.font.size == Pt(12), \
                            f"Executor font size is {run.font.size}, expected 12pt"
                    else:
                        assert run.font.size == Pt(14), \
                            f"Font size is {run.font.size} for text: {run.text[:50]}"

    def test_signature_table_center_aligned(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        sig_table = doc.tables[-1]
        # Right cell (name) should be center-aligned with firstLine indent
        right_cell = sig_table.rows[0].cells[1]
        for p in right_cell.paragraphs:
            if p.text.strip():
                assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER, \
                    f"Signature name should be center-aligned, got {p.alignment}"
                assert p.paragraph_format.first_line_indent is not None and \
                    p.paragraph_format.first_line_indent > 0, \
                    f"Signature firstLine should be set, got {p.paragraph_format.first_line_indent}"

    def test_appendix_numbered_items(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        # Find appendix table (contains "Приложение")
        appendix_table = None
        for table in doc.tables:
            if "Приложение" in table.rows[0].cells[0].text:
                appendix_table = table
                break
        assert appendix_table is not None

        right_cell = appendix_table.rows[0].cells[1]
        # Should have one paragraph per appendix item
        item_paras = [p for p in right_cell.paragraphs if p.text.strip()]
        assert len(item_paras) == len(sample_content["appendices"]), \
            f"Expected {len(sample_content['appendices'])} items, got {len(item_paras)}"

        # Each paragraph should start with "N. "
        for i, p in enumerate(item_paras):
            assert p.text.startswith(f"{i+1}. "), \
                f"Item {i+1} should start with '{i+1}. ', got: {p.text[:30]}"

    def test_no_leading_tabs_in_body(self, test_dir, config):
        from generate_docx import create_document

        content = {
            "title": "Test",
            "addressee": "Test",
            "resume": "\tПараграф с табом в начале.",
            "details": "",
            "conclusions": "",
            "signer_position": "Должность",
            "signer_name": "И.И. Иванов",
            "executor_name": "Иванов Иван Иванович",
            "executor_phone": "1234",
        }
        path = test_dir / "test_tab.docx"
        create_document(content, config, path)

        doc = Document(str(path))
        for p in doc.paragraphs:
            if p.text.strip():
                assert not p.text.startswith("\t"), \
                    f"Paragraph starts with tab: {p.text[:40]}"

    def test_all_tables_border_nil(self, test_dir, config, sample_content):
        from generate_docx import create_document

        path = test_dir / "test_v1.docx"
        create_document(sample_content, config, path)

        doc = Document(str(path))
        for idx, table in enumerate(doc.tables):
            # Check table-level borders
            tblPr = table._tbl.tblPr
            if tblPr is not None:
                borders = tblPr.find(qn("w:tblBorders"))
                if borders is not None:
                    for child in borders:
                        val = child.get(qn("w:val"))
                        assert val in ("none", "nil", None), \
                            f"Table {idx} border {child.tag} has val={val}"
            # Check cell-level borders
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.find(qn("w:tcPr"))
                    if tcPr is not None:
                        cb = tcPr.find(qn("w:tcBorders"))
                        if cb is not None:
                            for b in cb:
                                val = b.get(qn("w:val"))
                                assert val in ("none", "nil", None), \
                                    f"Table {idx} cell border {b.tag} has val={val}"

    def test_no_appendix_when_empty(self, test_dir, config, sample_content):
        from generate_docx import create_document

        content = dict(sample_content)
        content["appendices"] = []

        path = test_dir / "test_no_appendix.docx"
        create_document(content, config, path)

        doc = Document(str(path))
        # Should have 2 tables: header + signature (no appendix)
        assert len(doc.tables) == 2


# ---------------------------------------------------------------------------
# Tests: document patching
# ---------------------------------------------------------------------------

class TestPatchDocument:
    def test_patch_creates_new_version(self, test_dir, config, sample_content):
        from generate_docx import create_document, patch_document

        source = test_dir / "report_v1.docx"
        create_document(sample_content, config, source)

        fixes = {"fix_formatting": True}
        result = patch_document(source, fixes, config)
        assert result.exists()
        assert "v2" in result.name

    def test_patch_replaces_paragraph(self, test_dir, config, sample_content):
        from generate_docx import create_document, patch_document

        source = test_dir / "report_v1.docx"
        create_document(sample_content, config, source)

        doc = Document(str(source))
        # Find a non-empty paragraph index
        target_idx = None
        for i, p in enumerate(doc.paragraphs):
            if "нарушения" in p.text:
                target_idx = i
                break

        if target_idx is not None:
            fixes = {
                "replace_paragraphs": {
                    str(target_idx): "Обнаружены серьёзные нарушения."
                }
            }
            result = patch_document(source, fixes, config)
            doc2 = Document(str(result))
            assert "серьёзные" in doc2.paragraphs[target_idx].text

    def test_patch_fix_header_table(self, test_dir, config, sample_content):
        from generate_docx import create_document, patch_document

        source = test_dir / "report_v1.docx"
        create_document(sample_content, config, source)

        fixes = {"fix_header_table": True}
        result = patch_document(source, fixes, config)
        doc = Document(str(result))

        # Verify grid was set on header table
        grid = doc.tables[0]._tbl.find(qn("w:tblGrid"))
        assert grid is not None
        cols = grid.findall(qn("w:gridCol"))
        assert len(cols) == 2


# ---------------------------------------------------------------------------
# Tests: finalization
# ---------------------------------------------------------------------------

class TestFinalize:
    def test_finalize_without_password(self, test_dir, config, sample_content):
        """Finalize should handle missing password gracefully."""
        from generate_docx import create_document, finalize_document
        import os

        source = test_dir / "report_v1.docx"
        create_document(sample_content, config, source)

        # Remove password env var
        old_pwd = os.environ.pop("DOC_DEFAULT_PASSWORD", None)
        try:
            result = finalize_document(source, config)
            # Should still return docx path (possibly unencrypted)
            assert result["docx_path"] is not None
        finally:
            if old_pwd:
                os.environ["DOC_DEFAULT_PASSWORD"] = old_pwd
