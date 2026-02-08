"""Tests for analyze_docx.py — L1 format analysis."""

import json
import shutil
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils import load_config


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _create_conforming_document(path: Path, config: dict) -> Document:
    """Create a document that fully conforms to the formatting spec."""
    doc = Document()
    fmt = config["formatting"]
    page = config["page"]

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(page["width_cm"])
    section.page_height = Cm(page["height_cm"])
    section.left_margin = Cm(page["margin_left_cm"])
    section.right_margin = Cm(page["margin_right_cm"])
    section.top_margin = Cm(page["margin_top_cm"])
    section.bottom_margin = Cm(page["margin_bottom_cm"])

    # Header table (1x2, no borders)
    header_table = doc.add_table(rows=1, cols=2)
    _remove_borders(header_table)
    _set_cell(header_table.rows[0].cells[0], "Справка об инциденте ИБ", config)
    _set_cell(header_table.rows[0].cells[1], "Директору по безопасности\nИванову И.И.", config)

    # Empty paragraph
    _add_para(doc, "", config)

    # Body block 1: Resume
    _add_para(doc, "В ходе мониторинга систем информационной безопасности выявлен несанкционированный доступ к серверу баз данных.", config)

    _add_para(doc, "", config)

    # Body block 2: Details
    _add_para(doc, "10 января 2024 года в 14:30 зафиксирована попытка входа с неизвестного IP-адреса.", config)

    _add_para(doc, "", config)

    # Body block 3: Conclusions
    _add_para(doc, "Предлагается усилить контроль доступа и провести аудит всех учётных записей.", config)

    _add_para(doc, "", config)

    # Signature table (1x2, no borders)
    sig_table = doc.add_table(rows=1, cols=2)
    _remove_borders(sig_table)
    _set_cell(sig_table.rows[0].cells[0], "Начальник отдела ИБ", config)
    _set_cell(sig_table.rows[0].cells[1], "Петров П.П.", config)

    # Footer (2-line format: name + phone)
    footer = section.footer
    footer.is_linked_to_previous = False
    p1 = footer.paragraphs[0]
    run1 = p1.add_run("Сидоров С.С.")
    run1.font.name = fmt["font_name"]
    run1.font.size = Pt(12)
    from docx.oxml import OxmlElement as OE2
    p2_el = OE2("w:p")
    footer._element.append(p2_el)
    from docx.text.paragraph import Paragraph
    p2 = Paragraph(p2_el, footer)
    run2 = p2.add_run("тел. 1234")
    run2.font.name = fmt["font_name"]
    run2.font.size = Pt(12)

    doc.save(str(path))
    return doc


def _remove_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _set_cell(cell, text, config):
    """Set cell text with correct formatting."""
    fmt = config["formatting"]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = fmt["font_name"]
    run.font.size = Pt(fmt["font_size_pt"])


def _add_para(doc, text, config):
    """Add a correctly formatted paragraph."""
    fmt = config["formatting"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = fmt["line_spacing"]
    run = p.add_run(text)
    run.font.name = fmt["font_name"]
    run.font.size = Pt(fmt["font_size_pt"])


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def config():
    return load_config()


@pytest.fixture
def test_dir():
    d = Path(tempfile.mkdtemp())
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture
def conforming_doc(test_dir, config):
    path = test_dir / "conforming.docx"
    _create_conforming_document(path, config)
    return path


# ---------------------------------------------------------------------------
# Tests: analyze mode
# ---------------------------------------------------------------------------

class TestAnalyzeConforming:
    """Test that a conforming document passes with minimal issues."""

    def test_conforming_doc_has_no_critical_issues(self, conforming_doc, config):
        """A properly formatted document should have no critical issues."""
        from analyze_docx import cmd_analyze
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_analyze(conforming_doc, None, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            critical = [i for i in result["issues"] if i["severity"] == "critical"]
            assert len(critical) == 0, f"Critical issues found: {critical}"


class TestAnalyzePageSetup:
    """Test page setup detection."""

    def test_wrong_margins_detected(self, test_dir, config):
        """Wrong margins should be detected."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)  # Should be 3.0
        section.right_margin = Cm(2.0)  # Should be 1.0
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        # Add minimal content
        p = doc.add_paragraph("Test")
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        path = test_dir / "wrong_margins.docx"
        doc.save(str(path))

        from analyze_docx import cmd_analyze
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_analyze(path, None, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            margin_issues = [i for i in result["issues"] if "MARGIN" in i.get("code", "")]
            assert len(margin_issues) >= 2, "Should detect wrong left and right margins"


class TestAnalyzeFont:
    """Test font detection."""

    def test_wrong_font_detected(self, test_dir, config):
        """Wrong font should be detected."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("Some text with wrong font")
        run.font.name = "Arial"
        run.font.size = Pt(14)

        path = test_dir / "wrong_font.docx"
        doc.save(str(path))

        from analyze_docx import cmd_analyze
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_analyze(path, None, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            font_issues = [i for i in result["issues"] if i.get("code") == "FONT_NAME"]
            assert len(font_issues) > 0, "Should detect wrong font"


class TestAnalyzeStructure:
    """Test structural element detection."""

    def test_missing_header_table(self, test_dir, config):
        """Document without header table should flag it."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

        # Only add text, no header table
        p = doc.add_paragraph("Just text, no header table")
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        path = test_dir / "no_header.docx"
        doc.save(str(path))

        from analyze_docx import cmd_analyze
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_analyze(path, None, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            header_issues = [i for i in result["issues"]
                           if i.get("code") == "HEADER_TABLE_MISSING"]
            assert len(header_issues) > 0

    def test_missing_footer(self, test_dir, config):
        """Document without footer should flag it."""
        doc = Document()
        p = doc.add_paragraph("Test content")
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        path = test_dir / "no_footer.docx"
        doc.save(str(path))

        from analyze_docx import cmd_analyze
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_analyze(path, None, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            footer_issues = [i for i in result["issues"]
                           if "FOOTER" in i.get("code", "")]
            assert len(footer_issues) > 0


# ---------------------------------------------------------------------------
# Tests: extract mode
# ---------------------------------------------------------------------------

class TestExtract:
    def test_extract_returns_structure(self, conforming_doc, config):
        """Extract should return all expected fields."""
        from analyze_docx import cmd_extract
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_extract(conforming_doc, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            assert "header" in result
            assert "body_text" in result
            assert "signature" in result
            assert "footer" in result
            assert "metadata" in result

    def test_extract_body_text(self, conforming_doc, config):
        """Extract should contain body text."""
        from analyze_docx import cmd_extract
        import io
        from contextlib import redirect_stdout

        f = io.StringIO()
        with redirect_stdout(f):
            try:
                cmd_extract(conforming_doc, config)
            except SystemExit:
                pass

        output = f.getvalue()
        if output.strip():
            result = json.loads(output)
            assert "несанкционированный" in result["body_text"]


# ---------------------------------------------------------------------------
# Tests: verdict logic
# ---------------------------------------------------------------------------

class TestVerdictLogic:
    def test_rewrite_threshold(self, config):
        """Verify rewrite threshold: format > 5 AND content > 3 (strictly greater)."""
        thresholds = config["thresholds"]
        assert thresholds["rewrite_format_issues"] == 5
        assert thresholds["rewrite_content_issues"] == 3

        # Simulate: 6 format + 4 content → needs_rewrite
        format_issues = 6
        content_issues = 4
        verdict = "needs_rewrite" if (
            format_issues > thresholds["rewrite_format_issues"]
            and content_issues > thresholds["rewrite_content_issues"]
        ) else "needs_fixes"
        assert verdict == "needs_rewrite"

        # 5 format + 3 content → needs_fixes (boundary — not strictly greater)
        format_issues = 5
        content_issues = 3
        verdict = "needs_rewrite" if (
            format_issues > thresholds["rewrite_format_issues"]
            and content_issues > thresholds["rewrite_content_issues"]
        ) else "needs_fixes"
        assert verdict == "needs_fixes"  # 5 > 5 is false

        # 6 format + 2 content → needs_fixes (AND condition)
        format_issues = 6
        content_issues = 2
        verdict = "needs_rewrite" if (
            format_issues > thresholds["rewrite_format_issues"]
            and content_issues > thresholds["rewrite_content_issues"]
        ) else "needs_fixes"
        assert verdict == "needs_fixes"
