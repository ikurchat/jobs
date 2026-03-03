"""Document generation, patching, PDF conversion and encryption for doc-review skill.

CLI modes:
    python generate_docx.py create --input <content.json>
    python generate_docx.py patch --source <file.docx> --fixes <fixes.json>
    python generate_docx.py finalize --source <file.docx>

All operations happen in /dev/shm. Output is JSON on stdout.
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import shutil

from utils import (
    load_config,
    create_work_dir,
    cleanup_work_dir,
    encrypt_docx,
    encrypt_pdf,
    convert_to_pdf,
    next_version_path,
    sanitize_filename,
    output_json,
    output_error,
)


# ---------------------------------------------------------------------------
# Document formatting helpers
# ---------------------------------------------------------------------------

def _apply_paragraph_format(paragraph, config: dict) -> None:
    """Apply standard formatting to a paragraph from config."""
    fmt = config["formatting"]
    pf = paragraph.paragraph_format

    # Alignment
    alignment_map = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    pf.alignment = alignment_map.get(fmt["alignment"], WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Line spacing
    pf.line_spacing = fmt["line_spacing"]

    # First line indent
    if fmt.get("first_line_indent_cm"):
        pf.first_line_indent = Cm(fmt["first_line_indent_cm"])

    # Space before/after
    pf.space_before = Pt(fmt.get("space_before_pt", 0))
    pf.space_after = Pt(fmt.get("space_after_pt", 0))


def _apply_run_format(run, config: dict, bold: bool = False, italic: bool = False,
                      font_size_pt: int | None = None) -> None:
    """Apply standard font formatting to a run from config."""
    fmt = config["formatting"]
    size = font_size_pt or fmt["font_size_pt"]
    run.font.name = fmt["font_name"]
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

    # Ensure Times New Roman works for Cyrillic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), fmt["font_name"])
    rFonts.set(qn("w:hAnsi"), fmt["font_name"])
    rFonts.set(qn("w:cs"), fmt["font_name"])
    rFonts.set(qn("w:eastAsia"), fmt["font_name"])


def _apply_nbsp(text: str, config: dict) -> str:
    """Replace regular spaces with non-breaking spaces where needed."""
    nbsp = "\u00A0"
    nbsp_cfg = config.get("nbsp_rules", {})

    # Prepositions: "в ", "на ", etc. at word boundary → "в\u00A0"
    prepositions = nbsp_cfg.get("prepositions", [])
    if prepositions:
        pattern = r'\b(' + '|'.join(re.escape(p) for p in prepositions) + r')\s+'
        text = re.sub(pattern, lambda m: m.group(1) + nbsp, text, flags=re.IGNORECASE)

    # Patterns like №, ст. followed by space+digit
    for pat in nbsp_cfg.get("patterns", []):
        text = text.replace(f"{pat} ", f"{pat}{nbsp}")

    # Initials before surname: "И.О. " → "И.О.\u00A0"
    if nbsp_cfg.get("initials_before_surname"):
        text = re.sub(r'([А-ЯЁ]\.[А-ЯЁ]\.)\s+', r'\1' + nbsp, text)

    return text


def _add_formatted_paragraph(doc, text: str, config: dict, bold: bool = False,
                             italic: bool = False, font_size_pt: int | None = None) -> None:
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    _apply_paragraph_format(p, config)
    run = p.add_run(_apply_nbsp(text, config) if text else text)
    _apply_run_format(run, config, bold=bold, italic=italic, font_size_pt=font_size_pt)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _remove_table_borders(table) -> None:
    """Remove all visible borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)

    # Remove existing borders element if any
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Also remove cell borders
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            cell_borders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                cell_borders.append(border)
            existing_cb = tcPr.find(qn("w:tcBorders"))
            if existing_cb is not None:
                tcPr.remove(existing_cb)
            tcPr.append(cell_borders)


def _set_cell_text(cell, text: str, config: dict, bold: bool = False,
                   italic: bool = False, font_size_pt: int | None = None) -> None:
    """Set text in a table cell with standard formatting."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _apply_paragraph_format(p, config)
    run = p.add_run(_apply_nbsp(text, config) if text else text)
    _apply_run_format(run, config, bold=bold, italic=italic, font_size_pt=font_size_pt)


def _set_table_grid(table, col_widths_twips: list[int]) -> None:
    """Set tblGrid with explicit gridCol widths and full-width tblW.

    This is the authoritative way to control column widths in OOXML.
    Word uses tblGrid/gridCol as the ground truth; cell.width and tblW
    alone are often ignored.

    Args:
        table: python-docx Table object.
        col_widths_twips: List of column widths in twips (1/1440 inch).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    total_twips = sum(col_widths_twips)

    # --- tblW: total table width ---
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")

    # --- Remove tblLayout fixed (let Word auto-fit) ---
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is not None:
        tblPr.remove(tblLayout)

    # --- tblGrid: column definitions ---
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)

    new_grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        new_grid.append(col)

    # tblGrid must come right after tblPr in the XML
    tblPr_index = list(tbl).index(tblPr)
    tbl.insert(tblPr_index + 1, new_grid)

    # --- Cell widths (tcW) must match gridCol ---
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_twips):
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.insert(0, tcW)
                tcW.set(qn("w:w"), str(col_widths_twips[i]))
                tcW.set(qn("w:type"), "dxa")


def _page_width_twips(config: dict) -> int:
    """Return usable page width (between margins) in twips."""
    page = config["page"]
    width_cm = page["width_cm"] - page["margin_left_cm"] - page["margin_right_cm"]
    return int(width_cm / 2.54 * 1440)


def _create_1x2_table(doc, left_text: str, right_text: str, config: dict,
                       left_bold: bool = False, right_bold: bool = False,
                       left_align: str = "left", right_align: str = "left",
                       left_pct: float = 0.5,
                       left_font_size_pt: int | None = None,
                       left_col_twips: int | None = None):
    """Create a 1x2 table without borders spanning full page width.

    Args:
        left_align/right_align: "left", "right", "center", "justify".
        left_pct: Left column share (0..1). Ignored if left_col_twips set.
        left_font_size_pt: Override font size for left cell.
        left_col_twips: Exact left column width in twips (for auto-fit).
    Returns the created table.
    """
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _remove_table_borders(table)

    # Set proper grid with explicit column widths
    total_twips = _page_width_twips(config)
    if left_col_twips is not None:
        left_twips = left_col_twips
    else:
        left_twips = int(total_twips * left_pct)
    right_twips = total_twips - left_twips
    _set_table_grid(table, [left_twips, right_twips])

    _set_cell_text(table.rows[0].cells[0], left_text, config,
                   bold=left_bold, font_size_pt=left_font_size_pt)
    _set_cell_text(table.rows[0].cells[1], right_text, config, bold=right_bold)

    # Apply alignments
    for p in table.rows[0].cells[0].paragraphs:
        p.alignment = align_map.get(left_align, WD_ALIGN_PARAGRAPH.LEFT)
    for p in table.rows[0].cells[1].paragraphs:
        p.alignment = align_map.get(right_align, WD_ALIGN_PARAGRAPH.LEFT)

    return table


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _setup_page(doc, config: dict) -> None:
    """Configure page size and margins from config."""
    page = config["page"]
    section = doc.sections[0]

    section.page_width = Cm(page["width_cm"])
    section.page_height = Cm(page["height_cm"])
    section.left_margin = Cm(page["margin_left_cm"])
    section.right_margin = Cm(page["margin_right_cm"])
    section.top_margin = Cm(page["margin_top_cm"])
    section.bottom_margin = Cm(page["margin_bottom_cm"])


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _make_footer_run(paragraph, text: str, config: dict) -> None:
    """Add a run to a footer paragraph with proper formatting."""
    footer_cfg = config["structure"]["footer"]
    font_name = config["formatting"]["font_name"]
    font_size = footer_cfg.get("font_size_pt", 12)

    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Set Cyrillic font faces
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def _apply_footer_paragraph_format(paragraph) -> None:
    """Apply zero-spacing single-line format to a footer paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _clear_footer_xml(footer) -> None:
    """Properly clear footer by removing all child elements from XML."""
    footer_element = footer._element
    for child in list(footer_element):
        footer_element.remove(child)


def _clear_header_xml(header) -> None:
    """Clear header XML completely, removing SDT elements and adding empty paragraph."""
    header_element = header._element
    for child in list(header_element):
        header_element.remove(child)
    # Add one empty paragraph (required by OOXML)
    empty_p = OxmlElement("w:p")
    header_element.append(empty_p)


def _remove_title_page_flag(section) -> None:
    """Remove <w:titlePg/> from sectPr to disable 'Different First Page'."""
    sectPr = section._sectPr
    title_pg = sectPr.find(qn("w:titlePg"))
    if title_pg is not None:
        sectPr.remove(title_pg)


def _setup_footer(doc, executor_name: str, executor_phone: str, config: dict) -> None:
    """Add footer with executor info: name on line 1, phone on line 2, empty line after."""
    section = doc.sections[-1]  # Use last section (for multi-section docs)

    # Remove "Different First Page" flag so footer shows on page 1
    _remove_title_page_flag(section)

    # Clear any existing header SDT/PAGE elements (bug 5)
    header = section.header
    header.is_linked_to_previous = False
    _clear_header_xml(header)

    # Properly clear existing footer XML (bug 4)
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_footer_xml(footer)

    # Line 1: executor name
    p1 = footer.add_paragraph()
    _apply_footer_paragraph_format(p1)
    _make_footer_run(p1, executor_name, config)

    # Line 2: phone number
    p2 = footer.add_paragraph()
    _apply_footer_paragraph_format(p2)
    _make_footer_run(p2, executor_phone, config)

    # Line 3: empty line after executor info
    p3 = footer.add_paragraph()
    _apply_footer_paragraph_format(p3)


# ---------------------------------------------------------------------------
# Document creation
# ---------------------------------------------------------------------------

def create_document(content: dict, config: dict, output_path: Path) -> Path:
    """Create a new .docx document from structured content.

    Args:
        content: Dict with keys:
            - doc_type: str — "spravka" (default) or "sz"
            - title: str — document title
            - addressee: str — "Должность\\nФИО"
            - resume/intro: str — first block text
            - details: str — details block text
            - conclusions/conclusion: str — last block text
            - appendices: list[str] — optional list of appendix items
            - signer_position: str — signer's position
            - signer_name: str — signer's full name
            - executor_name: str — executor's name for footer
            - executor_phone: str — executor's phone for footer
        config: Parsed config.json.
        output_path: Where to save the document.

    Returns:
        Path to created file.
    """
    doc = Document()
    _setup_page(doc, config)

    # Determine document type
    doc_type = content.get("doc_type", "spravka")
    doc_types_cfg = config.get("doc_types", {})
    dt_cfg = doc_types_cfg.get(doc_type, {})
    conclusions_italic = dt_cfg.get("conclusions_style") == "italic"
    title_font_size = config["formatting"].get("title_font_size_pt", 12)

    # --- Header table (bold=False ALWAYS, title 12pt) ---
    title = content.get("title", "")
    addressee = content.get("addressee", "")
    _create_1x2_table(doc, title, addressee, config,
                       left_bold=False, right_bold=False,
                       left_align="left", right_align="right",
                       left_pct=0.6, left_font_size_pt=title_font_size)

    # --- 2 empty lines after header ---
    _add_formatted_paragraph(doc, "", config)
    _add_formatted_paragraph(doc, "", config)

    # --- Body: three blocks (NO separator paragraphs between them) ---
    # Block 1: Resume / Intro
    first_block_text = content.get("resume", "") or content.get("intro", "")
    if first_block_text:
        for para_text in first_block_text.split("\n\n"):
            if para_text.strip():
                _add_formatted_paragraph(doc, para_text.strip(), config)

    # Block 2: Details
    details_text = content.get("details", "")
    if details_text:
        for para_text in details_text.split("\n\n"):
            if para_text.strip():
                _add_formatted_paragraph(doc, para_text.strip(), config)

    # Block 3: Conclusions / Conclusion
    conclusions_text = content.get("conclusions", "") or content.get("conclusion", "")
    if conclusions_text:
        for para_text in conclusions_text.split("\n\n"):
            if para_text.strip():
                _add_formatted_paragraph(doc, para_text.strip(), config, italic=conclusions_italic)

    # --- Appendix table (optional, left col auto-fit to "Приложение:") ---
    appendices = content.get("appendices", [])
    if appendices:
        _add_formatted_paragraph(doc, "", config)
        appendix_label = config["structure"]["appendix"]["left_cell_text"]
        numbered_list = "\n".join(f"{i+1}. {item}" for i, item in enumerate(appendices))
        # Auto-fit left column: ~1600 twips fits "Приложение:" in TNR 14pt
        _create_1x2_table(doc, appendix_label, numbered_list, config,
                           left_col_twips=1600)

    # --- Signature table (left=LEFT, right=RIGHT) ---
    _add_formatted_paragraph(doc, "", config)
    signer_position = content.get("signer_position", "")
    signer_name = content.get("signer_name", "")
    _create_1x2_table(doc, signer_position, signer_name, config,
                       left_align="left", right_align="right",
                       left_pct=0.62)

    # --- Footer (executor info) ---
    executor_name = content.get("executor_name", "")
    # Auto-lookup phone from config executors directory
    executor_phone = content.get("executor_phone", "")
    if executor_name and not executor_phone:
        executors = config.get("executors", {})
        executor_phone = executors.get(executor_name, "")
    if executor_name:
        _setup_footer(doc, executor_name, executor_phone, config)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Document patching
# ---------------------------------------------------------------------------

def patch_document(source_path: Path, fixes: dict, config: dict,
                   output_path: Path | None = None) -> Path:
    """Apply fixes to an existing document.

    Args:
        source_path: Path to source .docx.
        fixes: Dict with fix instructions:
            - fix_formatting: bool — reapply all font/margin/spacing
            - fix_structure: bool — rebuild structural tables
            - replace_paragraphs: dict[str, str] — index -> new text
            - add_paragraphs: list[{"after": int, "text": str}]
            - remove_paragraphs: list[int] — indices to remove
        config: Parsed config.json.
        output_path: Where to save. Defaults to next version.

    Returns:
        Path to the patched document.
    """
    out_path = output_path or next_version_path(source_path)
    doc = Document(str(source_path))

    # Fix formatting
    if fixes.get("fix_formatting"):
        _fix_all_formatting(doc, config)

    # Fix page setup
    if fixes.get("fix_page_setup"):
        _setup_page(doc, config)

    # Fix header table grid (bug 12)
    if fixes.get("fix_header_table") and doc.tables:
        _fix_header_table_grid(doc.tables[0], config)

    # Replace paragraphs
    replacements = fixes.get("replace_paragraphs", {})
    for idx_str, new_text in replacements.items():
        idx = int(idx_str)
        if 0 <= idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            # Clear existing runs and add new one with proper formatting
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_text
                _apply_run_format(p.runs[0], config)
            else:
                run = p.add_run(new_text)
                _apply_run_format(run, config)

    doc.save(str(out_path))
    return out_path


def _fix_header_table_grid(table, config: dict) -> None:
    """Fix header table: set full-width grid (60/40), left=LEFT, right=RIGHT, no bold."""
    if len(table.columns) != 2:
        return
    total_twips = _page_width_twips(config)
    left_twips = int(total_twips * 0.6)
    right_twips = total_twips - left_twips
    _set_table_grid(table, [left_twips, right_twips])

    # Left cell: LEFT alignment, no bold
    for p in table.rows[0].cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = False

    # Right cell: RIGHT alignment, no bold
    for p in table.rows[0].cells[1].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.bold = False


def _fix_all_formatting(doc, config: dict) -> None:
    """Reapply correct formatting to all paragraphs and runs."""
    for paragraph in doc.paragraphs:
        _apply_paragraph_format(paragraph, config)
        for run in paragraph.runs:
            _apply_run_format(run, config)

    # Fix tables (text formatting + grid)
    for i, table in enumerate(doc.tables):
        # Fix header table grid (first table)
        if i == 0 and len(table.columns) == 2:
            _fix_header_table_grid(table, config)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_paragraph_format(paragraph, config)
                    for run in paragraph.runs:
                        _apply_run_format(run, config)


# ---------------------------------------------------------------------------
# Finalization (encrypt + PDF)
# ---------------------------------------------------------------------------

def finalize_document(source_path: Path, config: dict) -> dict:
    """Finalize a document: encrypt .docx, convert to PDF, encrypt PDF.

    Returns:
        {
            "docx_path": str,
            "pdf_path": str | null,
            "encrypted": true
        }
    """
    result = {
        "docx_path": None,
        "pdf_path": None,
        "encrypted": False,
    }

    # Step 1: Encrypt .docx
    encrypted_docx = source_path.parent / f"final_{source_path.name}"
    try:
        shutil.copy2(str(source_path), str(encrypted_docx))
        encrypt_docx(encrypted_docx)
        result["docx_path"] = str(encrypted_docx)
        result["encrypted"] = True
    except Exception as e:
        # If encryption fails, use unencrypted version
        result["docx_path"] = str(source_path)
        result["encryption_error"] = str(e)

    # Step 2: Convert to PDF
    pdf_path = convert_to_pdf(source_path, source_path.parent)
    if pdf_path:
        # Step 3: Encrypt PDF
        encrypted_pdf = source_path.parent / f"final_{pdf_path.name}"
        try:
            shutil.copy2(str(pdf_path), str(encrypted_pdf))
            encrypt_pdf(encrypted_pdf)
            result["pdf_path"] = str(encrypted_pdf)
            # Remove unencrypted PDF
            pdf_path.unlink(missing_ok=True)
        except Exception as e:
            result["pdf_path"] = str(pdf_path)
            result["pdf_encryption_error"] = str(e)
    else:
        result["pdf_path"] = None
        result["pdf_note"] = "LibreOffice не доступен, PDF не создан"

    # Step 4: Remove unencrypted source if encryption succeeded
    if result["encrypted"] and encrypted_docx.exists():
        source_path.unlink(missing_ok=True)

    return result


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate/edit .docx documents")
    subparsers = parser.add_subparsers(dest="mode", required=True)

    # create
    p_create = subparsers.add_parser("create", help="Create new document from JSON")
    p_create.add_argument("--input", type=Path, required=True,
                          help="JSON file with document content")
    p_create.add_argument("--output", type=Path, default=None,
                          help="Output .docx path")
    p_create.add_argument("--config", type=Path, default=None)

    # patch
    p_patch = subparsers.add_parser("patch", help="Apply fixes to existing document")
    p_patch.add_argument("--source", type=Path, required=True,
                         help="Source .docx file")
    p_patch.add_argument("--fixes", type=Path, required=True,
                         help="JSON file with fixes")
    p_patch.add_argument("--output", type=Path, default=None)
    p_patch.add_argument("--config", type=Path, default=None)

    # finalize
    p_final = subparsers.add_parser("finalize", help="Encrypt + PDF + encrypt PDF")
    p_final.add_argument("--source", type=Path, required=True,
                         help="Source .docx file to finalize")
    p_final.add_argument("--config", type=Path, default=None)

    args = parser.parse_args()

    try:
        config = load_config(args.config)
    except Exception as e:
        output_error(f"Config error: {e}", code=2)

    work_dir = None
    try:
        if args.mode == "create":
            with open(args.input, "r", encoding="utf-8") as f:
                content = json.load(f)

            if not args.output:
                work_dir = create_work_dir(config)
                title = content.get("title", "document")
                safe_name = sanitize_filename(title)
                out_path = work_dir / f"{safe_name}_v1.docx"
            else:
                out_path = args.output

            result_path = create_document(content, config, out_path)
            output_json({
                "mode": "create",
                "output_path": str(result_path),
                "status": "ok",
            })

        elif args.mode == "patch":
            with open(args.fixes, "r", encoding="utf-8") as f:
                fixes = json.load(f)

            result_path = patch_document(
                args.source, fixes, config, args.output
            )
            output_json({
                "mode": "patch",
                "source_path": str(args.source),
                "output_path": str(result_path),
                "status": "ok",
            })

        elif args.mode == "finalize":
            result = finalize_document(args.source, config)
            result["mode"] = "finalize"
            result["status"] = "ok"
            output_json(result)

    except Exception as e:
        output_error(str(e))
    finally:
        if work_dir:
            cleanup_work_dir(work_dir)


if __name__ == "__main__":
    main()
