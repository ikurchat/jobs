"""
Скрипт заполнения еженедельного отчёта ЦОК.

Использование:
    python fill_weekly_report.py fill --template /path/to/plan.docx --tasks /path/to/tasks.json --output /path/to/report.docx
    python fill_weekly_report.py add_additional --doc /path/to/report.docx --tasks /path/to/unplanned.json --output /path/to/report_final.docx
    python fill_weekly_report.py validate --doc /path/to/report.docx
"""

from __future__ import annotations

import argparse
import copy
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE = Pt(12)

# Table indices
HEADER_TABLE_IDX = 0
DATA_TABLE_IDX = 1

# Data rows start at index 2 (R0=section header, R1=sub-header)
DATA_ROW_START = 2

# Column indices
COL_NUM = 0
COL_ACTIVITY = 1
COL_DATES = 2
COL_RESPONSIBLE = 3
COL_MARK = 4

NUM_COLUMNS = 5


# ---------------------------------------------------------------------------
# load_template
# ---------------------------------------------------------------------------

def load_template(path: str) -> Document:
    """Load .docx and verify structure: 2 tables, 5 columns, landscape."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    doc = Document(path)

    # Check landscape orientation
    section = doc.sections[0]
    if section.orientation != WD_ORIENT.LANDSCAPE:
        raise ValueError(
            f"Expected landscape orientation, got "
            f"{'portrait' if section.orientation == WD_ORIENT.PORTRAIT else 'unknown'}"
        )

    # Check tables
    if len(doc.tables) < 2:
        raise ValueError(f"Expected at least 2 tables, found {len(doc.tables)}")

    header_table = doc.tables[HEADER_TABLE_IDX]
    data_table = doc.tables[DATA_TABLE_IDX]

    # Check columns
    if len(header_table.columns) != NUM_COLUMNS:
        raise ValueError(
            f"Header table: expected {NUM_COLUMNS} columns, found {len(header_table.columns)}"
        )
    if len(data_table.columns) != NUM_COLUMNS:
        raise ValueError(
            f"Data table: expected {NUM_COLUMNS} columns, found {len(data_table.columns)}"
        )

    return doc


# ---------------------------------------------------------------------------
# detect_font_settings
# ---------------------------------------------------------------------------

def detect_font_settings(doc: Document) -> dict:
    """Detect font settings from existing data cells (columns 1-3)."""
    data_table = doc.tables[DATA_TABLE_IDX]
    font_names: list[str] = []
    font_sizes: list[int] = []

    for row in data_table.rows[DATA_ROW_START:]:
        for col_idx in (COL_ACTIVITY, COL_DATES, COL_RESPONSIBLE):
            cell = row.cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.name:
                        font_names.append(run.font.name)
                    if run.font.size is not None:
                        font_sizes.append(run.font.size)

    # Determine most common values
    detected_name = DEFAULT_FONT_NAME
    detected_size = DEFAULT_FONT_SIZE

    if font_names:
        detected_name = max(set(font_names), key=font_names.count)
    if font_sizes:
        detected_size = max(set(font_sizes), key=font_sizes.count)

    return {"font_name": detected_name, "font_size": detected_size}


# ---------------------------------------------------------------------------
# fill_marks
# ---------------------------------------------------------------------------

def fill_marks(
    doc: Document,
    tasks: list[dict],
    font_settings: dict,
) -> Document:
    """Fill column 4 (marks) for each data row based on matching tasks."""
    data_table = doc.tables[DATA_TABLE_IDX]
    font_name = font_settings["font_name"]
    font_size = font_settings["font_size"]

    for row in data_table.rows[DATA_ROW_START:]:
        activity_text = row.cells[COL_ACTIVITY].text.strip()
        if not activity_text:
            continue

        matched = _match_tasks(activity_text, tasks)
        mark_text = _format_mark(matched, activity_text)

        _set_cell_text(row.cells[COL_MARK], mark_text, font_name, font_size)

    return doc


def _match_tasks(activity_text: str, tasks: list[dict]) -> list[dict]:
    """Find tasks matching an activity row by plan_item_hint or keyword overlap."""
    activity_lower = activity_text.lower()
    activity_words = set(activity_lower.split())

    matched: list[dict] = []
    for task in tasks:
        if task.get("is_unplanned"):
            continue

        # Priority 1: plan_item_hint
        hint = (task.get("plan_item_hint") or "").lower()
        if hint and hint in activity_lower:
            matched.append(task)
            continue

        # Priority 2: keyword overlap (title words)
        title_words = set(task.get("title", "").lower().split())
        # Exclude short common words
        significant = {w for w in title_words if len(w) > 3}
        overlap = significant & activity_words
        if len(overlap) >= 2:
            matched.append(task)

    return matched


def _format_mark(matched: list[dict], activity_text: str) -> str:
    """Format mark text from matched tasks."""
    if not matched:
        short = activity_text[:60]
        return f"[Ожидается уточнение: статус — {short}]"

    parts: list[str] = []
    for task in matched:
        status = task.get("status", "pending")
        result = task.get("result", "")
        notes = task.get("notes", "")
        ref = task.get("regulatory_ref", "")

        if status == "done":
            line = "Выполнено."
            if result:
                line += f" {result}."
            if ref:
                line += f" {ref}."
        elif status == "in_progress":
            line = "В работе."
            detail = result or notes
            if detail:
                line += f" {detail}."
        elif status == "cancelled":
            line = "Отменено."
            if notes:
                line += f" {notes}."
        else:
            line = f"[Ожидается уточнение: {result or notes or 'статус не определён'}]"

        parts.append(line.strip())

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# add_additional_section
# ---------------------------------------------------------------------------

def add_additional_section(
    doc: Document,
    unplanned_tasks: list[dict],
    font_settings: dict,
) -> Document:
    """Add 'Дополнительные мероприятия' section after last data row."""
    # Filter: only done + is_unplanned
    done_unplanned = [
        t for t in unplanned_tasks
        if t.get("is_unplanned") and t.get("status") == "done"
    ]

    if not done_unplanned:
        return doc

    data_table = doc.tables[DATA_TABLE_IDX]
    font_name = font_settings["font_name"]
    font_size = font_settings["font_size"]

    # Add merged header row "Дополнительные мероприятия"
    _add_merged_row(data_table, "Дополнительные мероприятия", font_name, Pt(12), bold=True)

    # Add data rows
    for i, task in enumerate(done_unplanned, 1):
        title = task.get("title", "")
        deadline = task.get("deadline", "")
        result = task.get("result", "")
        ref = task.get("regulatory_ref", "")

        mark = "Выполнено."
        if result:
            mark += f" {result}."
        if ref:
            mark += f" {ref}."

        _add_data_row(
            data_table,
            num=str(i),
            activity=title,
            dates=deadline,
            responsible="",
            mark=mark.strip(),
            font_name=font_name,
            font_size=font_size,
        )

    return doc


def _add_merged_row(
    table, text: str, font_name: str, font_size, bold: bool = True
) -> None:
    """Add a horizontally merged row spanning all 5 columns."""
    row = table.add_row()
    # Merge all cells
    cell = row.cells[0]
    for i in range(1, NUM_COLUMNS):
        cell.merge(row.cells[i])

    # Set text and formatting
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold

    # Center alignment
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_data_row(
    table,
    num: str,
    activity: str,
    dates: str,
    responsible: str,
    mark: str,
    font_name: str,
    font_size,
) -> None:
    """Add a regular data row with 5 columns."""
    row = table.add_row()
    values = [num, activity, dates, responsible, mark]
    for ci, text in enumerate(values):
        _set_cell_text(row.cells[ci], text, font_name, font_size)


# ---------------------------------------------------------------------------
# change_title
# ---------------------------------------------------------------------------

def change_title(doc: Document, new_title: str) -> Document:
    """Replace P6 title text, preserving formatting (14pt TNR Bold Center)."""
    if len(doc.paragraphs) < 7:
        raise ValueError(f"Document has only {len(doc.paragraphs)} paragraphs, expected >= 7")

    p6 = doc.paragraphs[6]

    # Preserve formatting from first run
    old_runs = p6.runs
    ref_font_name = DEFAULT_FONT_NAME
    ref_font_size = Pt(14)
    ref_bold = True

    if old_runs:
        r0 = old_runs[0]
        if r0.font.name:
            ref_font_name = r0.font.name
        if r0.font.size is not None:
            ref_font_size = r0.font.size
        if r0.font.bold is not None:
            ref_bold = r0.font.bold

    # Clear existing runs
    for run in old_runs:
        run._element.getparent().remove(run._element)

    # Add new run
    new_run = p6.add_run(new_title)
    new_run.font.name = ref_font_name
    new_run.font.size = ref_font_size
    new_run.font.bold = ref_bold

    return doc


# ---------------------------------------------------------------------------
# validate_formatting
# ---------------------------------------------------------------------------

def validate_formatting(doc: Document) -> list[str]:
    """Validate column 4 formatting and completeness."""
    issues: list[str] = []
    data_table = doc.tables[DATA_TABLE_IDX]
    font_settings = detect_font_settings(doc)
    expected_name = font_settings["font_name"]
    expected_size = font_settings["font_size"]

    for ri, row in enumerate(data_table.rows[DATA_ROW_START:], start=DATA_ROW_START):
        cell = row.cells[COL_MARK]
        activity_cell = row.cells[COL_ACTIVITY]
        activity_text = activity_cell.text.strip()

        # Skip merged/section rows
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None and int(gs.get(qn("w:val"), "1")) > 1:
                continue

        mark_text = cell.text.strip()

        # Check empty marks for non-empty activities
        if activity_text and not mark_text:
            issues.append(f"R{ri}: пустая отметка для «{activity_text[:50]}»")
            continue

        if not mark_text:
            continue

        # Check font on each run
        for pi, paragraph in enumerate(cell.paragraphs):
            for runi, run in enumerate(paragraph.runs):
                run_text = run.text.strip()
                if not run_text:
                    continue

                # Font name
                actual_name = run.font.name
                if actual_name and actual_name != expected_name:
                    issues.append(
                        f"R{ri}/C4: шрифт «{actual_name}» вместо «{expected_name}»"
                    )

                # Font size
                actual_size = run.font.size
                if actual_size is not None and actual_size != expected_size:
                    expected_pt = expected_size // 12700
                    actual_pt = actual_size // 12700
                    issues.append(
                        f"R{ri}/C4: размер {actual_pt}pt вместо {expected_pt}pt"
                    )

    return issues


# ---------------------------------------------------------------------------
# save_doc
# ---------------------------------------------------------------------------

def save_doc(doc: Document, output_path: str) -> None:
    """Save document, verifying column 4 fonts are correct."""
    font_settings = detect_font_settings(doc)
    data_table = doc.tables[DATA_TABLE_IDX]

    for row in data_table.rows[DATA_ROW_START:]:
        cell = row.cells[COL_MARK]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                if run.font.name is None:
                    run.font.name = font_settings["font_name"]
                if run.font.size is None:
                    run.font.size = font_settings["font_size"]

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str, font_name: str, font_size) -> None:
    """Clear cell and set text with explicit font on every run."""
    # Clear existing content
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    # Use first paragraph
    p = cell.paragraphs[0]
    # Split by newlines to handle multi-line text
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = font_size


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def cmd_fill(args: argparse.Namespace) -> None:
    """Fill marks from tasks JSON."""
    doc = load_template(args.template)

    with open(args.tasks, "r", encoding="utf-8") as f:
        tasks = json.load(f)

    font_settings = detect_font_settings(doc)
    print(f"Font: {font_settings['font_name']}, size: {font_settings['font_size'] // 12700}pt")

    if args.title:
        change_title(doc, args.title)

    fill_marks(doc, tasks, font_settings)

    issues = validate_formatting(doc)
    if issues:
        print(f"\nWarnings ({len(issues)}):")
        for issue in issues:
            print(f"  - {issue}")

    save_doc(doc, args.output)
    print(f"\nSaved: {args.output}")


def cmd_add_additional(args: argparse.Namespace) -> None:
    """Add additional (unplanned) section."""
    doc = load_template(args.doc)

    with open(args.tasks, "r", encoding="utf-8") as f:
        tasks = json.load(f)

    font_settings = detect_font_settings(doc)
    add_additional_section(doc, tasks, font_settings)

    save_doc(doc, args.output)
    print(f"Saved: {args.output}")


def cmd_validate(args: argparse.Namespace) -> None:
    """Validate document formatting."""
    doc = load_template(args.doc)
    font_settings = detect_font_settings(doc)

    print(f"Document: {args.doc}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Detected font: {font_settings['font_name']}, {font_settings['font_size'] // 12700}pt")

    data_table = doc.tables[DATA_TABLE_IDX]
    total_rows = len(data_table.rows)
    data_rows = total_rows - DATA_ROW_START
    print(f"Data rows: {data_rows} (R{DATA_ROW_START}-R{total_rows - 1})")

    # Count filled marks
    filled = 0
    empty = 0
    for row in data_table.rows[DATA_ROW_START:]:
        tc = row.cells[COL_MARK]._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None and int(gs.get(qn("w:val"), "1")) > 1:
                continue
        mark = row.cells[COL_MARK].text.strip()
        activity = row.cells[COL_ACTIVITY].text.strip()
        if activity:
            if mark:
                filled += 1
            else:
                empty += 1

    print(f"Filled marks: {filled}/{filled + empty}")

    # Title
    if len(doc.paragraphs) > 6:
        title = doc.paragraphs[6].text
        print(f"Title: {title}")

    # Validate
    issues = validate_formatting(doc)
    if issues:
        print(f"\nIssues ({len(issues)}):")
        for issue in issues:
            print(f"  ❌ {issue}")
        sys.exit(1)
    else:
        print("\n✅ All formatting checks passed")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Weekly report filler for ЦОК",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # fill
    p_fill = sub.add_parser("fill", help="Fill marks from tasks")
    p_fill.add_argument("--template", required=True, help="Template .docx path")
    p_fill.add_argument("--tasks", required=True, help="Tasks JSON path")
    p_fill.add_argument("--output", required=True, help="Output .docx path")
    p_fill.add_argument("--title", default=None, help="Override title (P6)")

    # add_additional
    p_add = sub.add_parser("add_additional", help="Add unplanned tasks section")
    p_add.add_argument("--doc", required=True, help="Report .docx path")
    p_add.add_argument("--tasks", required=True, help="Unplanned tasks JSON")
    p_add.add_argument("--output", required=True, help="Output .docx path")

    # validate
    p_val = sub.add_parser("validate", help="Validate formatting")
    p_val.add_argument("--doc", required=True, help="Document .docx path")

    args = parser.parse_args()

    if args.command == "fill":
        cmd_fill(args)
    elif args.command == "add_additional":
        cmd_add_additional(args)
    elif args.command == "validate":
        cmd_validate(args)


if __name__ == "__main__":
    main()
