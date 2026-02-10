#!/usr/bin/env python3
"""Level 1 (automated format) analysis of .docx documents.

CLI modes:
  analyze <file> [--prev <prev_file>]  - full formatting & structure check
  extract <file>                       - extract text for L2/L3 analysis

Outputs structured JSON to stdout.
"""

import argparse
import os
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from utils import (
    load_config,
    is_encrypted,
    decrypt_docx,
    create_work_dir,
    cleanup_work_dir,
    cm_to_emu,
    emu_to_cm,
    diff_documents,
    output_json,
    output_error,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

SEVERITY_CRITICAL = "critical"
SEVERITY_DESIRABLE = "desirable"
SEVERITY_RECOMMENDATION = "recommendation"

SEVERITY_MAP = {
    "PAGE_SIZE": SEVERITY_CRITICAL,
    "MARGIN_LEFT": SEVERITY_DESIRABLE,
    "MARGIN_RIGHT": SEVERITY_DESIRABLE,
    "MARGIN_TOP": SEVERITY_DESIRABLE,
    "MARGIN_BOTTOM": SEVERITY_DESIRABLE,
    "FONT_NAME": SEVERITY_DESIRABLE,
    "FONT_SIZE": SEVERITY_DESIRABLE,
    "ALIGNMENT": SEVERITY_DESIRABLE,
    "LINE_SPACING": SEVERITY_DESIRABLE,
    "FIRST_LINE_INDENT": SEVERITY_RECOMMENDATION,
    "HEADER_TABLE_MISSING": SEVERITY_CRITICAL,
    "HEADER_TABLE_STRUCTURE": SEVERITY_DESIRABLE,
    "HEADER_TABLE_BORDERS": SEVERITY_DESIRABLE,
    "BODY_SUBHEADERS": SEVERITY_RECOMMENDATION,
    "APPENDIX_TABLE_STRUCTURE": SEVERITY_DESIRABLE,
    "APPENDIX_TABLE_BORDERS": SEVERITY_DESIRABLE,
    "SIGNATURE_TABLE_MISSING": SEVERITY_DESIRABLE,
    "SIGNATURE_TABLE_STRUCTURE": SEVERITY_DESIRABLE,
    "SIGNATURE_TABLE_BORDERS": SEVERITY_DESIRABLE,
    "FOOTER_MISSING": SEVERITY_CRITICAL,
    "FOOTER_CONTENT": SEVERITY_DESIRABLE,
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _open_document(file_path: Path, config: dict) -> tuple[Document, Path | None, bool]:
    """Open a .docx, decrypting if necessary.

    Returns:
        (Document, work_dir_or_None, was_encrypted)
    """
    encrypted = is_encrypted(file_path)
    work_dir = None

    if encrypted:
        work_dir = create_work_dir(config)
        decrypted_path = work_dir / file_path.name
        decrypt_docx(file_path, decrypted_path)
        doc = Document(str(decrypted_path))
    else:
        doc = Document(str(file_path))

    return doc, work_dir, encrypted


def _check_file(file_path: Path, config: dict) -> None:
    """Validate that the file exists, is within size limits, and has a safe path."""
    resolved = file_path.resolve()
    if ".." in file_path.parts:
        output_error(f"Path traversal detected: {file_path}")
    if not resolved.exists():
        output_error(f"File not found: {file_path}")
    if not resolved.is_file():
        output_error(f"Not a file: {file_path}")

    max_size = config.get("analysis", {}).get("max_file_size_mb", 50)
    size_mb = file_path.stat().st_size / (1024 * 1024)
    if size_mb > max_size:
        output_error(f"File too large: {size_mb:.1f} MB (max {max_size} MB)")


def _table_has_visible_borders(table) -> bool:
    """Check if a table or its cells have visible borders.

    Inspects tblBorders on the table and tcBorders on each cell.
    A border is considered visible if its val is not 'none', 'nil', or absent.
    """
    visible_vals = lambda el: el is not None and el.get(qn("w:val")) not in (
        "none", "nil", None,
    )

    # Check table-level borders
    tbl_pr = table._tbl.tblPr
    if tbl_pr is not None:
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is not None:
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border_el = tbl_borders.find(qn(f"w:{side}"))
                if visible_vals(border_el):
                    return True

    # Check cell-level borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            if tc_pr is not None:
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is not None:
                    for side in ("top", "left", "bottom", "right"):
                        border_el = tc_borders.find(qn(f"w:{side}"))
                        if visible_vals(border_el):
                            return True

    return False


def _get_cell_text(table, row: int, col: int) -> str:
    """Safely get text from a table cell."""
    try:
        return table.cell(row, col).text.strip()
    except (IndexError, AttributeError):
        return ""


def _approx_eq_emu(actual_emu: int | None, expected_cm: float, tolerance_cm: float) -> bool:
    """Compare EMU value against expected cm within tolerance."""
    if actual_emu is None:
        return False
    actual_cm = emu_to_cm(actual_emu)
    return abs(actual_cm - expected_cm) <= tolerance_cm


def _approx_eq_pt(actual_pt: float | None, expected_pt: float, tolerance_pt: float) -> bool:
    """Compare point values within tolerance."""
    if actual_pt is None:
        return False
    return abs(actual_pt - expected_pt) <= tolerance_pt


def _resolve_font_name(run) -> str | None:
    """Resolve effective font name for a run.

    python-docx returns None when the value is inherited from the style chain.
    We walk: run.font -> run style -> paragraph style -> document defaults.
    """
    # Direct run font
    if run.font.name is not None:
        return run.font.name

    # Check rPr XML directly for ascii font
    rpr = run._r.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                val = rfonts.get(qn(attr))
                if val:
                    return val

    # Try the run's style
    if run.style and run.style.font and run.style.font.name:
        return run.style.font.name

    return None


def _resolve_font_size_pt(run) -> float | None:
    """Resolve effective font size for a run in points."""
    if run.font.size is not None:
        return run.font.size.pt

    # Check rPr XML directly
    rpr = run._r.find(qn("w:rPr"))
    if rpr is not None:
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            val = sz.get(qn("w:val"))
            if val is not None:
                try:
                    return int(val) / 2.0
                except ValueError:
                    pass

    # Try the run's style
    if run.style and run.style.font and run.style.font.size:
        return run.style.font.size.pt

    return None


def _resolve_line_spacing(paragraph) -> float | None:
    """Resolve line spacing as a float multiplier (e.g. 1.15).

    python-docx stores line spacing in multiple ways:
    - As a float multiplier when line_spacing_rule is MULTIPLE
    - As Pt when line_spacing_rule is EXACTLY or AT_LEAST
    - As raw XML value in 240ths of a line
    """
    pf = paragraph.paragraph_format

    if pf.line_spacing is not None:
        # If it is a float (MULTIPLE rule), return directly
        if isinstance(pf.line_spacing, (int, float)):
            return float(pf.line_spacing)
        # If it is Pt-like (Emu internally), check the rule
        if hasattr(pf.line_spacing, "pt"):
            return pf.line_spacing.pt

    # Try XML directly
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is not None:
        spacing = ppr.find(qn("w:spacing"))
        if spacing is not None:
            line_val = spacing.get(qn("w:line"))
            line_rule = spacing.get(qn("w:lineRule"))
            if line_val is not None:
                try:
                    val = int(line_val)
                    # 'auto' or absent rule means value is in 240ths of a line
                    if line_rule in (None, "auto"):
                        return val / 240.0
                    # 'exact' or 'atLeast' means value is in twips (1/20 pt)
                    return val / 20.0
                except ValueError:
                    pass

    return None


def _resolve_alignment(paragraph):
    """Resolve paragraph alignment, considering style inheritance."""
    if paragraph.alignment is not None:
        return paragraph.alignment

    # Check pPr XML
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is not None:
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            val = jc.get(qn("w:val"))
            jc_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            return jc_map.get(val)

    # Try paragraph style
    if paragraph.style and paragraph.style.paragraph_format:
        return paragraph.style.paragraph_format.alignment

    return None


def _resolve_first_line_indent_emu(paragraph) -> int | None:
    """Resolve first-line indent in EMU."""
    pf = paragraph.paragraph_format
    if pf.first_line_indent is not None:
        return int(pf.first_line_indent)

    # Check pPr XML
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is not None:
        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            first_line = ind.get(qn("w:firstLine"))
            if first_line is not None:
                try:
                    # Value is in twips; 1 twip = 1/20 pt = 1/1440 inch = 635 EMU
                    return int(first_line) * 635
                except ValueError:
                    pass

    return None


def _is_body_paragraph(paragraph) -> bool:
    """Determine if a paragraph is a body text paragraph (not empty, not in table)."""
    return bool(paragraph.text.strip())


def _has_bold_or_underline(paragraph) -> bool:
    """Check if a paragraph is formatted as a visible subheader (bold and/or underline)."""
    text = paragraph.text.strip()
    if not text:
        return False

    # All runs must be bold/underlined for it to be a "subheader"
    runs_with_text = [r for r in paragraph.runs if r.text.strip()]
    if not runs_with_text:
        return False

    all_bold = all(r.bold for r in runs_with_text)
    all_underline = all(r.underline for r in runs_with_text)

    return all_bold or all_underline


# ---------------------------------------------------------------------------
# Structural analysis helpers
# ---------------------------------------------------------------------------

def _classify_document_elements(doc: Document) -> dict:
    """Walk document body and classify top-level elements.

    Returns dict with:
      tables: list of (index_in_body, table)
      paragraphs: list of (index_in_body, paragraph)
      body_elements_order: list of ('table', idx) or ('paragraph', idx)
    """
    tables = []
    paragraphs = []
    order = []

    table_idx = 0
    para_idx = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "tbl":
            if table_idx < len(doc.tables):
                tables.append((len(order), doc.tables[table_idx]))
                order.append(("table", table_idx))
                table_idx += 1
        elif tag == "p":
            if para_idx < len(doc.paragraphs):
                paragraphs.append((len(order), doc.paragraphs[para_idx]))
                order.append(("paragraph", para_idx))
                para_idx += 1

    return {
        "tables": tables,
        "paragraphs": paragraphs,
        "order": order,
    }


def _identify_header_table(tables: list, config: dict) -> Any | None:
    """Identify the header table (first table, should be 1-row x 2-col)."""
    header_cfg = config.get("structure", {}).get("header", {})
    expected_rows = header_cfg.get("rows", 1)
    expected_cols = header_cfg.get("cols", 2)

    if not tables:
        return None

    # Header table is expected to be the first table
    _, first_table = tables[0]
    if len(first_table.rows) >= expected_rows and len(first_table.columns) >= expected_cols:
        return first_table
    return None


def _identify_appendix_table(tables: list, config: dict) -> Any | None:
    """Identify the appendix table by looking for 'Приложение:' text."""
    appendix_cfg = config.get("structure", {}).get("appendix", {})
    marker_text = appendix_cfg.get("left_cell_text", "Приложение:")

    for _, table in tables:
        try:
            cell_text = _get_cell_text(table, 0, 0)
            if marker_text.lower() in cell_text.lower():
                return table
        except Exception:
            continue
    return None


def _identify_signature_table(tables: list, config: dict, appendix_table) -> Any | None:
    """Identify the signature table (last non-appendix table, or second-to-last if appendix exists)."""
    if len(tables) < 2:
        # If there's only the header table, no signature table
        return None

    # Walk from the end; skip appendix table if present
    for _, table in reversed(tables):
        if appendix_table is not None and table._tbl is appendix_table._tbl:
            continue
        # The header table is tables[0]; skip it
        if tables and table._tbl is tables[0][1]._tbl:
            continue
        return table

    return None


def _identify_body_paragraphs(doc: Document, elements: dict, header_table, appendix_table, signature_table) -> list:
    """Identify body text paragraphs: those between header and appendix/signature tables.

    Returns list of (original_paragraph_index, paragraph).
    """
    # Determine positional boundaries in the element order
    header_pos = -1
    appendix_pos = len(elements["order"])
    signature_pos = len(elements["order"])

    for pos, (etype, eidx) in enumerate(elements["order"]):
        if etype == "table":
            table = doc.tables[eidx]
            if header_table is not None and table._tbl is header_table._tbl:
                header_pos = pos
            if appendix_table is not None and table._tbl is appendix_table._tbl:
                appendix_pos = min(appendix_pos, pos)
            if signature_table is not None and table._tbl is signature_table._tbl:
                signature_pos = min(signature_pos, pos)

    end_pos = min(appendix_pos, signature_pos)

    # Collect body paragraphs with 1-based global paragraph numbering
    result = []
    para_global = 0
    for pos, (etype, eidx) in enumerate(elements["order"]):
        if etype == "paragraph":
            para_global += 1
            if header_pos < pos < end_pos:
                result.append((para_global, doc.paragraphs[eidx]))

    return result


# ---------------------------------------------------------------------------
# Checkers
# ---------------------------------------------------------------------------

def _check_page_setup(doc: Document, config: dict, tolerance_cm: float) -> list[dict]:
    """Check page size and margins against config."""
    issues = []
    page_cfg = config.get("page", {})

    for section in doc.sections:
        # Page size
        expected_w = page_cfg.get("width_cm", 21.0)
        expected_h = page_cfg.get("height_cm", 29.7)

        if not _approx_eq_emu(section.page_width, expected_w, tolerance_cm):
            actual_w = emu_to_cm(section.page_width) if section.page_width else 0
            issues.append({
                "level": "format",
                "severity": SEVERITY_MAP["PAGE_SIZE"],
                "code": "PAGE_SIZE",
                "message": f"Размер страницы: ширина {actual_w:.1f} см вместо {expected_w} см",
                "location": "section",
                "expected": f"{expected_w}x{expected_h} cm",
                "actual": f"{actual_w:.1f} cm width",
            })

        if not _approx_eq_emu(section.page_height, expected_h, tolerance_cm):
            actual_h = emu_to_cm(section.page_height) if section.page_height else 0
            issues.append({
                "level": "format",
                "severity": SEVERITY_MAP["PAGE_SIZE"],
                "code": "PAGE_SIZE",
                "message": f"Размер страницы: высота {actual_h:.1f} см вместо {expected_h} см",
                "location": "section",
                "expected": f"{expected_w}x{expected_h} cm",
                "actual": f"{actual_h:.1f} cm height",
            })

        # Margins
        margin_checks = [
            ("MARGIN_LEFT", "margin_left_cm", section.left_margin, "левое"),
            ("MARGIN_RIGHT", "margin_right_cm", section.right_margin, "правое"),
            ("MARGIN_TOP", "margin_top_cm", section.top_margin, "верхнее"),
            ("MARGIN_BOTTOM", "margin_bottom_cm", section.bottom_margin, "нижнее"),
        ]

        for code, cfg_key, actual_emu, label in margin_checks:
            expected_cm = page_cfg.get(cfg_key, 2.0)
            if not _approx_eq_emu(actual_emu, expected_cm, tolerance_cm):
                actual_cm = emu_to_cm(actual_emu) if actual_emu else 0
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP[code],
                    "code": code,
                    "message": f"Поле {label}: {actual_cm:.2f} см вместо {expected_cm} см",
                    "location": "section",
                    "expected": f"{expected_cm} cm",
                    "actual": f"{actual_cm:.2f} cm",
                })

        # Only check first section (typical for single-section docs)
        break

    return issues


def _check_paragraph_formatting(
    body_paragraphs: list,
    config: dict,
    tolerance_pt: float,
    tolerance_cm: float,
) -> list[dict]:
    """Check font, size, alignment, spacing, indent for body paragraphs."""
    issues = []
    fmt_cfg = config.get("formatting", {})
    expected_font = fmt_cfg.get("font_name", "Times New Roman")
    expected_size = fmt_cfg.get("font_size_pt", 14)
    expected_alignment_str = fmt_cfg.get("alignment", "justify")
    expected_alignment = ALIGNMENT_MAP.get(expected_alignment_str)
    expected_spacing = fmt_cfg.get("line_spacing", 1.15)
    expected_indent_cm = fmt_cfg.get("first_line_indent_cm")

    # Track which paragraphs already reported for each issue type to avoid spam
    reported_fonts = set()
    reported_sizes = set()

    for para_num, paragraph in body_paragraphs:
        if not _is_body_paragraph(paragraph):
            continue

        location = f"paragraph:{para_num}"

        # --- Font name and size (check each run) ---
        for run in paragraph.runs:
            if not run.text.strip():
                continue

            # Font name
            if para_num not in reported_fonts:
                font_name = _resolve_font_name(run)
                if font_name is not None and font_name != expected_font:
                    issues.append({
                        "level": "format",
                        "severity": SEVERITY_MAP["FONT_NAME"],
                        "code": "FONT_NAME",
                        "message": f"Шрифт абзаца {para_num}: {font_name} вместо {expected_font}",
                        "location": location,
                        "expected": expected_font,
                        "actual": font_name,
                    })
                    reported_fonts.add(para_num)

            # Font size
            if para_num not in reported_sizes:
                font_size = _resolve_font_size_pt(run)
                if font_size is not None and not _approx_eq_pt(font_size, expected_size, tolerance_pt):
                    issues.append({
                        "level": "format",
                        "severity": SEVERITY_MAP["FONT_SIZE"],
                        "code": "FONT_SIZE",
                        "message": f"Размер шрифта абзаца {para_num}: {font_size}pt вместо {expected_size}pt",
                        "location": location,
                        "expected": f"{expected_size}pt",
                        "actual": f"{font_size}pt",
                    })
                    reported_sizes.add(para_num)

        # --- Alignment ---
        actual_alignment = _resolve_alignment(paragraph)
        if actual_alignment is not None and expected_alignment is not None:
            if actual_alignment != expected_alignment:
                align_names = {
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                }
                actual_name = align_names.get(actual_alignment, str(actual_alignment))
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["ALIGNMENT"],
                    "code": "ALIGNMENT",
                    "message": f"Выравнивание абзаца {para_num}: {actual_name} вместо {expected_alignment_str}",
                    "location": location,
                    "expected": expected_alignment_str,
                    "actual": actual_name,
                })

        # --- Line spacing ---
        actual_spacing = _resolve_line_spacing(paragraph)
        if actual_spacing is not None:
            # Only flag if it looks like a multiplier (reasonable range 0.5-5.0)
            if 0.5 <= actual_spacing <= 5.0:
                if not _approx_eq_pt(actual_spacing, expected_spacing, 0.02):
                    issues.append({
                        "level": "format",
                        "severity": SEVERITY_MAP["LINE_SPACING"],
                        "code": "LINE_SPACING",
                        "message": f"Межстрочный интервал абзаца {para_num}: {actual_spacing:.2f} вместо {expected_spacing}",
                        "location": location,
                        "expected": str(expected_spacing),
                        "actual": f"{actual_spacing:.2f}",
                    })

        # --- First line indent ---
        if expected_indent_cm is not None:
            actual_indent_emu = _resolve_first_line_indent_emu(paragraph)
            if actual_indent_emu is not None:
                if not _approx_eq_emu(actual_indent_emu, expected_indent_cm, tolerance_cm):
                    actual_indent_cm = emu_to_cm(actual_indent_emu)
                    issues.append({
                        "level": "format",
                        "severity": SEVERITY_MAP["FIRST_LINE_INDENT"],
                        "code": "FIRST_LINE_INDENT",
                        "message": f"Отступ первой строки абзаца {para_num}: {actual_indent_cm:.2f} см вместо {expected_indent_cm} см",
                        "location": location,
                        "expected": f"{expected_indent_cm} cm",
                        "actual": f"{actual_indent_cm:.2f} cm",
                    })

    return issues


def _check_structure(doc: Document, config: dict) -> tuple[list[dict], dict, list, Any, Any, Any]:
    """Check structural elements: header table, body, appendix, signature, footer.

    Returns:
        (issues, structure_info, body_paragraphs, header_table, appendix_table, signature_table)
    """
    issues = []
    struct_cfg = config.get("structure", {})

    elements = _classify_document_elements(doc)
    tables = elements["tables"]

    # --- Header table ---
    header_table = _identify_header_table(tables, config)
    has_header_table = header_table is not None

    if not has_header_table:
        issues.append({
            "level": "format",
            "severity": SEVERITY_MAP["HEADER_TABLE_MISSING"],
            "code": "HEADER_TABLE_MISSING",
            "message": "Отсутствует таблица-шапка документа",
            "location": "document",
            "expected": "1x2 table at document start",
            "actual": "missing",
        })
    else:
        header_cfg = struct_cfg.get("header", {})
        expected_rows = header_cfg.get("rows", 1)
        expected_cols = header_cfg.get("cols", 2)

        actual_rows = len(header_table.rows)
        actual_cols = len(header_table.columns)
        if actual_rows != expected_rows or actual_cols != expected_cols:
            issues.append({
                "level": "format",
                "severity": SEVERITY_MAP["HEADER_TABLE_STRUCTURE"],
                "code": "HEADER_TABLE_STRUCTURE",
                "message": f"Таблица-шапка: {actual_rows}x{actual_cols} вместо {expected_rows}x{expected_cols}",
                "location": "header_table",
                "expected": f"{expected_rows}x{expected_cols}",
                "actual": f"{actual_rows}x{actual_cols}",
            })

        if not header_cfg.get("borders", False):
            if _table_has_visible_borders(header_table):
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["HEADER_TABLE_BORDERS"],
                    "code": "HEADER_TABLE_BORDERS",
                    "message": "Таблица-шапка имеет видимые границы",
                    "location": "header_table",
                    "expected": "no visible borders",
                    "actual": "borders present",
                })

    # --- Appendix table ---
    appendix_table = _identify_appendix_table(tables, config)
    has_appendix_table = appendix_table is not None

    if has_appendix_table:
        appendix_cfg = struct_cfg.get("appendix", {})
        expected_rows = appendix_cfg.get("rows", 1)
        expected_cols = appendix_cfg.get("cols", 2)

        actual_rows = len(appendix_table.rows)
        actual_cols = len(appendix_table.columns)
        if actual_rows != expected_rows or actual_cols != expected_cols:
            issues.append({
                "level": "format",
                "severity": SEVERITY_MAP["APPENDIX_TABLE_STRUCTURE"],
                "code": "APPENDIX_TABLE_STRUCTURE",
                "message": f"Таблица приложений: {actual_rows}x{actual_cols} вместо {expected_rows}x{expected_cols}",
                "location": "appendix_table",
                "expected": f"{expected_rows}x{expected_cols}",
                "actual": f"{actual_rows}x{actual_cols}",
            })

        if not appendix_cfg.get("borders", False):
            if _table_has_visible_borders(appendix_table):
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["APPENDIX_TABLE_BORDERS"],
                    "code": "APPENDIX_TABLE_BORDERS",
                    "message": "Таблица приложений имеет видимые границы",
                    "location": "appendix_table",
                    "expected": "no visible borders",
                    "actual": "borders present",
                })

    # --- Signature table ---
    signature_table = _identify_signature_table(tables, config, appendix_table)
    has_signature_table = signature_table is not None

    if not has_signature_table:
        issues.append({
            "level": "format",
            "severity": SEVERITY_MAP["SIGNATURE_TABLE_MISSING"],
            "code": "SIGNATURE_TABLE_MISSING",
            "message": "Отсутствует таблица подписи",
            "location": "document",
            "expected": "1x2 table near document end",
            "actual": "missing",
        })
    else:
        sig_cfg = struct_cfg.get("signature", {})
        expected_rows = sig_cfg.get("rows", 1)
        expected_cols = sig_cfg.get("cols", 2)

        actual_rows = len(signature_table.rows)
        actual_cols = len(signature_table.columns)
        if actual_rows != expected_rows or actual_cols != expected_cols:
            issues.append({
                "level": "format",
                "severity": SEVERITY_MAP["SIGNATURE_TABLE_STRUCTURE"],
                "code": "SIGNATURE_TABLE_STRUCTURE",
                "message": f"Таблица подписи: {actual_rows}x{actual_cols} вместо {expected_rows}x{expected_cols}",
                "location": "signature_table",
                "expected": f"{expected_rows}x{expected_cols}",
                "actual": f"{actual_rows}x{actual_cols}",
            })

        if not sig_cfg.get("borders", False):
            if _table_has_visible_borders(signature_table):
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["SIGNATURE_TABLE_BORDERS"],
                    "code": "SIGNATURE_TABLE_BORDERS",
                    "message": "Таблица подписи имеет видимые границы",
                    "location": "signature_table",
                    "expected": "no visible borders",
                    "actual": "borders present",
                })

    # --- Footer ---
    has_footer = False
    footer_text = ""
    for section in doc.sections:
        try:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                footer_paras = footer.paragraphs
                if footer_paras:
                    footer_text = " ".join(p.text.strip() for p in footer_paras if p.text.strip())
                    if footer_text:
                        has_footer = True
            elif footer:
                # Linked to previous -- still may have content
                footer_paras = footer.paragraphs
                if footer_paras:
                    footer_text = " ".join(p.text.strip() for p in footer_paras if p.text.strip())
                    if footer_text:
                        has_footer = True
        except Exception:
            pass
        break  # Check first section only

    if not has_footer:
        issues.append({
            "level": "format",
            "severity": SEVERITY_MAP["FOOTER_MISSING"],
            "code": "FOOTER_MISSING",
            "message": "Отсутствует колонтитул с информацией об исполнителе",
            "location": "footer",
            "expected": "footer with executor info",
            "actual": "missing",
        })
    else:
        # Check footer content: should have at least 2 non-empty lines (name + phone)
        # or a single line with "тел." format
        footer_cfg = struct_cfg.get("footer", {})
        footer_format = footer_cfg.get("format", "two_lines")
        footer_lines = [p.text.strip() for p in doc.sections[0].footer.paragraphs if p.text.strip()]

        if footer_format == "two_lines":
            # Two-line format: name on line 1, phone on line 2
            if len(footer_lines) < 2:
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["FOOTER_CONTENT"],
                    "code": "FOOTER_CONTENT",
                    "message": f"Колонтитул: ожидается 2 строки (ФИО + телефон), найдено {len(footer_lines)}",
                    "location": "footer",
                    "expected": "2 lines: name + phone",
                    "actual": footer_text[:100],
                })
        else:
            # Legacy single-line format with "тел."
            template = footer_cfg.get("content_template", "{executor_name}, тел. {executor_phone}")
            if "тел." in template and "тел." not in footer_text.lower():
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["FOOTER_CONTENT"],
                    "code": "FOOTER_CONTENT",
                    "message": f"Колонтитул не соответствует шаблону: '{footer_text[:80]}'",
                    "location": "footer",
                    "expected": template,
                    "actual": footer_text[:100],
                })

    # --- Body subheaders ---
    body_paragraphs = _identify_body_paragraphs(
        doc, elements, header_table, appendix_table, signature_table,
    )

    if not struct_cfg.get("body", {}).get("visible_subheaders", False):
        for para_num, paragraph in body_paragraphs:
            if _has_bold_or_underline(paragraph):
                issues.append({
                    "level": "format",
                    "severity": SEVERITY_MAP["BODY_SUBHEADERS"],
                    "code": "BODY_SUBHEADERS",
                    "message": f"Абзац {para_num}: обнаружен видимый подзаголовок (жирный/подчёркнутый текст)",
                    "location": f"paragraph:{para_num}",
                    "expected": "no visible subheaders",
                    "actual": paragraph.text[:60],
                })

    structure_info = {
        "has_header_table": has_header_table,
        "has_appendix_table": has_appendix_table,
        "has_signature_table": has_signature_table,
        "has_footer": has_footer,
        "body_paragraph_count": len([
            (n, p) for n, p in body_paragraphs if _is_body_paragraph(p)
        ]),
        "table_count": len(doc.tables),
    }

    return issues, structure_info, body_paragraphs, header_table, appendix_table, signature_table


# ---------------------------------------------------------------------------
# Mode: analyze
# ---------------------------------------------------------------------------

def cmd_analyze(file_path: Path, prev_path: Path | None, config: dict) -> None:
    """Run full L1 analysis and output JSON."""
    _check_file(file_path, config)
    if prev_path:
        _check_file(prev_path, config)

    analysis_cfg = config.get("analysis", {})
    tolerance_pt = analysis_cfg.get("tolerance_pt", 0.5)
    tolerance_cm = analysis_cfg.get("tolerance_cm", 0.1)
    thresholds = config.get("thresholds", {})

    work_dir = None
    try:
        doc, work_dir, encrypted = _open_document(file_path, config)

        all_issues: list[dict] = []

        # 1. Page setup
        all_issues.extend(_check_page_setup(doc, config, tolerance_cm))

        # 2. Structural checks (also returns body paragraphs for formatting checks)
        struct_issues, structure_info, body_paragraphs, *_ = _check_structure(doc, config)
        all_issues.extend(struct_issues)

        # 3. Paragraph formatting
        all_issues.extend(
            _check_paragraph_formatting(body_paragraphs, config, tolerance_pt, tolerance_cm)
        )

        # 4. Diff with previous version
        diff_result = None
        if prev_path:
            try:
                # If prev file is also encrypted, need to decrypt it too
                if is_encrypted(prev_path):
                    if work_dir is None:
                        work_dir = create_work_dir(config)
                    prev_decrypted = work_dir / f"prev_{prev_path.name}"
                    decrypt_docx(prev_path, prev_decrypted)
                    # Use decrypted current if needed
                    current_for_diff = (work_dir / file_path.name) if encrypted else file_path
                    diff_result = diff_documents(prev_decrypted, current_for_diff)
                else:
                    current_for_diff = (work_dir / file_path.name) if encrypted else file_path
                    diff_result = diff_documents(prev_path, current_for_diff)
            except Exception as e:
                diff_result = {"error": str(e)}

        # 5. Build summary
        total = len(all_issues)
        critical_count = sum(1 for i in all_issues if i["severity"] == SEVERITY_CRITICAL)
        desirable_count = sum(1 for i in all_issues if i["severity"] == SEVERITY_DESIRABLE)
        recommendation_count = sum(1 for i in all_issues if i["severity"] == SEVERITY_RECOMMENDATION)
        format_issues = sum(1 for i in all_issues if i["level"] == "format")
        content_issues = sum(1 for i in all_issues if i["level"] == "content")

        # 6. Determine verdict
        if total == 0:
            verdict = "ok"
        elif (
            format_issues > thresholds.get("rewrite_format_issues", 5)
            and content_issues > thresholds.get("rewrite_content_issues", 3)
        ):
            verdict = "needs_rewrite"
        else:
            verdict = "needs_fixes"

        rewrite_suggested = verdict == "needs_rewrite"

        result = {
            "file": file_path.name,
            "encrypted": encrypted,
            "issues": all_issues,
            "structure": structure_info,
            "summary": {
                "total_issues": total,
                "critical": critical_count,
                "desirable": desirable_count,
                "recommendation": recommendation_count,
                "format_issues": format_issues,
                "content_issues": content_issues,
            },
            "verdict": verdict,
            "rewrite_suggested": rewrite_suggested,
            "diff": diff_result,
        }

        output_json(result)

    except SystemExit:
        raise
    except Exception as e:
        output_error(f"Analysis failed: {e}")
    finally:
        if work_dir:
            cleanup_work_dir(work_dir)


# ---------------------------------------------------------------------------
# Mode: extract
# ---------------------------------------------------------------------------

def cmd_extract(file_path: Path, config: dict) -> None:
    """Extract text content from the document for L2/L3 analysis."""
    _check_file(file_path, config)

    work_dir = None
    try:
        doc, work_dir, _encrypted = _open_document(file_path, config)

        elements = _classify_document_elements(doc)
        tables = elements["tables"]

        # Identify structural tables
        header_table = _identify_header_table(tables, config)
        appendix_table = _identify_appendix_table(tables, config)
        signature_table = _identify_signature_table(tables, config, appendix_table)

        # --- Header ---
        header_data = {"title": "", "addressee": ""}
        if header_table is not None:
            header_data["title"] = _get_cell_text(header_table, 0, 0)
            header_data["addressee"] = _get_cell_text(header_table, 0, 1)

        # --- Body text ---
        body_paragraphs = _identify_body_paragraphs(
            doc, elements, header_table, appendix_table, signature_table,
        )
        body_texts = []
        for _, paragraph in body_paragraphs:
            text = paragraph.text.strip()
            if text:
                body_texts.append(text)
        body_text = "\n".join(body_texts)

        # --- Appendix ---
        appendix_items = []
        if appendix_table is not None:
            # Right cell typically contains numbered list
            try:
                right_cell_text = _get_cell_text(appendix_table, 0, 1)
                if right_cell_text:
                    # Split by newlines for individual items
                    appendix_items = [
                        line.strip()
                        for line in right_cell_text.split("\n")
                        if line.strip()
                    ]
            except Exception:
                pass

        # --- Signature ---
        signature_data = {"position": "", "name": ""}
        if signature_table is not None:
            signature_data["position"] = _get_cell_text(signature_table, 0, 0)
            signature_data["name"] = _get_cell_text(signature_table, 0, 1)

        # --- Footer ---
        footer_text = ""
        for section in doc.sections:
            try:
                footer = section.footer
                if footer:
                    paras = footer.paragraphs
                    if paras:
                        footer_text = " ".join(
                            p.text.strip() for p in paras if p.text.strip()
                        )
            except Exception:
                pass
            break

        # --- Metadata ---
        metadata = {}
        try:
            props = doc.core_properties
            metadata["author"] = props.author or ""
            metadata["created"] = props.created.isoformat() if props.created else ""
            metadata["modified"] = props.modified.isoformat() if props.modified else ""
        except Exception:
            metadata = {"author": "", "created": "", "modified": ""}

        result = {
            "file": file_path.name,
            "header": header_data,
            "body_text": body_text,
            "appendix": appendix_items,
            "signature": signature_data,
            "footer": footer_text,
            "metadata": metadata,
        }

        output_json(result)

    except SystemExit:
        raise
    except Exception as e:
        output_error(f"Extraction failed: {e}")
    finally:
        if work_dir:
            cleanup_work_dir(work_dir)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Level 1 automated analysis of .docx documents",
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    # --- analyze ---
    analyze_parser = subparsers.add_parser(
        "analyze",
        help="Full L1 formatting and structure analysis",
    )
    analyze_parser.add_argument("file", type=str, help="Path to .docx file")
    analyze_parser.add_argument(
        "--prev", type=str, default=None,
        help="Path to previous version for diff",
    )

    # --- extract ---
    extract_parser = subparsers.add_parser(
        "extract",
        help="Extract text content for L2/L3 analysis",
    )
    extract_parser.add_argument("file", type=str, help="Path to .docx file")

    args = parser.parse_args()
    config = load_config()

    if args.command == "analyze":
        file_path = Path(args.file).resolve()
        prev_path = Path(args.prev).resolve() if args.prev else None
        cmd_analyze(file_path, prev_path, config)

    elif args.command == "extract":
        file_path = Path(args.file).resolve()
        cmd_extract(file_path, config)


if __name__ == "__main__":
    main()
