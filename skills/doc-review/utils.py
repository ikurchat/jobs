"""Shared utilities for doc-review skill.

Handles: config loading, password management, file operations in /dev/shm,
encryption/decryption, versioning, diff between document versions, cleanup.
"""

import json
import os
import re
import shutil
import sys
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SKILL_DIR = Path(__file__).parent
DEFAULT_CONFIG_PATH = SKILL_DIR / "config.json"


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

def load_config(config_path: Path | None = None) -> dict:
    """Load config.json and return parsed dict."""
    path = config_path or DEFAULT_CONFIG_PATH
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Unit conversions
# ---------------------------------------------------------------------------

def cm_to_emu(cm: float) -> int:
    """Convert centimeters to EMU (English Metric Units)."""
    return int(cm * 360000)


def pt_to_half_points(pt: float) -> int:
    """Convert points to half-points (used in some XML attributes)."""
    return int(pt * 2)


def emu_to_cm(emu: int) -> float:
    """Convert EMU to centimeters."""
    return emu / 360000


def half_points_to_pt(hp: int) -> float:
    """Convert half-points to points."""
    return hp / 2


# ---------------------------------------------------------------------------
# Password management
# ---------------------------------------------------------------------------

def get_password() -> str:
    """Read document password from environment variable.

    Returns password string.
    Raises RuntimeError if not set.
    """
    env_var = "DOC_DEFAULT_PASSWORD"
    password = os.environ.get(env_var)
    if not password:
        raise RuntimeError(
            f"Environment variable {env_var} is not set. "
            "Cannot encrypt/decrypt documents."
        )
    return password


# ---------------------------------------------------------------------------
# Working directory management (/dev/shm)
# ---------------------------------------------------------------------------

def create_work_dir(config: dict) -> Path:
    """Create a unique working directory in /dev/shm.

    Returns Path to the created directory.
    """
    base = Path(config.get("work_dir", "/dev/shm/doc-review"))
    session_id = uuid.uuid4().hex[:12]
    work_dir = base / session_id
    work_dir.mkdir(parents=True, exist_ok=False, mode=0o700)
    return work_dir


def cleanup_work_dir(work_dir: Path) -> None:
    """Remove working directory and all its contents.

    Safe to call even if directory doesn't exist.
    """
    if work_dir and work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)


def cleanup_all_work_dirs(config: dict) -> int:
    """Remove all working directories (startup cleanup).

    Returns number of directories cleaned.
    """
    base = Path(config.get("work_dir", "/dev/shm/doc-review"))
    if not base.exists():
        return 0
    count = 0
    for child in base.iterdir():
        if child.is_dir():
            shutil.rmtree(child, ignore_errors=True)
            count += 1
    return count


# ---------------------------------------------------------------------------
# Encryption / Decryption
# ---------------------------------------------------------------------------

def is_encrypted(file_path: Path) -> bool:
    """Check if a .docx file is password-protected (OLE2 encrypted)."""
    import msoffcrypto
    try:
        with open(file_path, "rb") as f:
            office_file = msoffcrypto.OfficeFile(f)
            return office_file.is_encrypted()
    except Exception:
        return False


def decrypt_docx(encrypted_path: Path, output_path: Path) -> Path:
    """Decrypt a password-protected .docx file.

    Args:
        encrypted_path: Path to encrypted .docx.
        output_path: Where to save decrypted file.

    Returns:
        Path to decrypted file.

    Raises:
        RuntimeError: If decryption fails.
    """
    import msoffcrypto
    password = get_password()
    try:
        with open(encrypted_path, "rb") as f_in:
            office_file = msoffcrypto.OfficeFile(f_in)
            office_file.load_key(password=password)
            with open(output_path, "wb") as f_out:
                office_file.decrypt(f_out)
    except Exception as e:
        raise RuntimeError(f"Failed to decrypt document: {e}")
    return output_path


def encrypt_docx(source_path: Path, output_path: Path | None = None) -> Path:
    """Password-protect a .docx file using msoffcrypto-tool.

    Args:
        source_path: Path to unencrypted .docx.
        output_path: Where to save. If None, overwrites source.

    Returns:
        Path to encrypted file.
    """
    import msoffcrypto
    password = get_password()
    out = output_path or source_path
    temp_path = source_path.with_suffix(".tmp_enc")

    try:
        with open(source_path, "rb") as f_in:
            office_file = msoffcrypto.OfficeFile(f_in)
            with open(temp_path, "wb") as f_out:
                office_file.save(f_out, password)
        shutil.move(str(temp_path), str(out))
    except Exception as e:
        if temp_path.exists():
            temp_path.unlink()
        raise RuntimeError(f"Failed to encrypt .docx: {e}")
    return out


def encrypt_pdf(source_path: Path, output_path: Path | None = None) -> Path:
    """Password-protect a .pdf file using pikepdf.

    Args:
        source_path: Path to unencrypted .pdf.
        output_path: Where to save. If None, overwrites source.

    Returns:
        Path to encrypted file.
    """
    import pikepdf
    password = get_password()
    out = output_path or source_path
    temp_path = source_path.with_suffix(".tmp_enc")

    try:
        with pikepdf.open(source_path) as pdf:
            pdf.save(
                str(temp_path),
                encryption=pikepdf.Encryption(
                    owner=password,
                    user=password,
                    R=6,
                )
            )
        shutil.move(str(temp_path), str(out))
    except Exception as e:
        if temp_path.exists():
            temp_path.unlink()
        raise RuntimeError(f"Failed to encrypt PDF: {e}")
    return out


# ---------------------------------------------------------------------------
# PDF conversion
# ---------------------------------------------------------------------------

def convert_to_pdf(docx_path: Path, output_dir: Path | None = None) -> Path | None:
    """Convert .docx to .pdf using LibreOffice headless.

    Returns Path to PDF or None if LibreOffice is not available.
    """
    import subprocess
    out_dir = output_dir or docx_path.parent

    try:
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--norestore",
                "--convert-to", "pdf",
                "--outdir", str(out_dir),
                str(docx_path),
            ],
            capture_output=True,
            timeout=120,
            cwd=str(out_dir),
        )
        if result.returncode != 0:
            return None
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return None

    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if pdf_path.exists():
        return pdf_path
    return None


# ---------------------------------------------------------------------------
# Versioning
# ---------------------------------------------------------------------------

_VERSION_RE = re.compile(r"^(.+?)_v(\d+)$")


def parse_version(filename: str) -> tuple[str, int]:
    """Parse filename into (base_name, version).

    Examples:
        "report_v1" -> ("report", 1)
        "report" -> ("report", 0)
    """
    stem = Path(filename).stem
    m = _VERSION_RE.match(stem)
    if m:
        return m.group(1), int(m.group(2))
    return stem, 0


def next_version_path(current_path: Path) -> Path:
    """Generate path for the next version of a document.

    Examples:
        report_v1.docx -> report_v2.docx
        report.docx -> report_v1.docx
    """
    base_name, version = parse_version(current_path.name)
    new_version = version + 1 if version > 0 else 1
    new_name = f"{base_name}_v{new_version}{current_path.suffix}"
    return current_path.parent / new_name


# ---------------------------------------------------------------------------
# Diff between document versions
# ---------------------------------------------------------------------------

def diff_documents(old_path: Path, new_path: Path) -> dict:
    """Compare two .docx files and return a structured diff.

    Returns:
        {
            "added_paragraphs": [...],
            "removed_paragraphs": [...],
            "changed_paragraphs": [{"old": ..., "new": ..., "index": ...}],
            "summary": "..."
        }
    """
    old_doc = Document(str(old_path))
    new_doc = Document(str(new_path))

    old_texts = [p.text.strip() for p in old_doc.paragraphs if p.text.strip()]
    new_texts = [p.text.strip() for p in new_doc.paragraphs if p.text.strip()]

    old_set = set(old_texts)
    new_set = set(new_texts)

    added = [t for t in new_texts if t not in old_set]
    removed = [t for t in old_texts if t not in new_set]

    # Simple change detection via LCS-like approach
    changed = []
    from difflib import SequenceMatcher
    matcher = SequenceMatcher(None, old_texts, new_texts)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "replace":
            for k in range(min(i2 - i1, j2 - j1)):
                changed.append({
                    "old": old_texts[i1 + k][:200],
                    "new": new_texts[j1 + k][:200],
                    "index": j1 + k,
                })

    summary_parts = []
    if added:
        summary_parts.append(f"добавлено {len(added)} абзацев")
    if removed:
        summary_parts.append(f"удалено {len(removed)} абзацев")
    if changed:
        summary_parts.append(f"изменено {len(changed)} абзацев")
    if not summary_parts:
        summary_parts.append("изменений не обнаружено")

    return {
        "added_paragraphs": [t[:200] for t in added[:20]],
        "removed_paragraphs": [t[:200] for t in removed[:20]],
        "changed_paragraphs": changed[:20],
        "summary": "; ".join(summary_parts),
    }


# ---------------------------------------------------------------------------
# Filename sanitization
# ---------------------------------------------------------------------------

def sanitize_filename(name: str) -> str:
    """Sanitize filename: remove path components, special chars."""
    # Take only the basename
    name = Path(name).name
    # Remove dangerous characters
    name = re.sub(r'[^\w\s\-\.]', '', name)
    # Collapse whitespace
    name = re.sub(r'\s+', '_', name.strip())
    if not name:
        name = "document"
    return name


# ---------------------------------------------------------------------------
# Safe JSON output
# ---------------------------------------------------------------------------

def output_json(data: dict) -> None:
    """Print JSON to stdout for consumption by SKILL.md."""
    print(json.dumps(data, ensure_ascii=False, indent=2))


def output_error(message: str, code: int = 1) -> None:
    """Print error JSON to stderr and exit."""
    print(json.dumps({"error": message}, ensure_ascii=False), file=sys.stderr)
    sys.exit(code)
